import os
import logging
import time
import threading
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
import shutil
import re

# Set up logger
logger = logging.getLogger(__name__)

class GitService:
    """Service for Git operations with thread-safe repository access"""
    
    def __init__(self, default_author_name: str = "OtherTales", default_author_email: str = "system@othertales.com"):
        """Initialize the Git service"""
        self.repo_locks = {}
        self.default_author_name = default_author_name
        self.default_author_email = default_author_email
        
        # Try importing gitpython - optional dependency
        try:
            import git
            self.git = git
        except ImportError:
            logger.error("GitPython is not installed. Git functionality will be limited.")
            self.git = None
            
    def _get_repo_lock(self, repo_path: str):
        """Get or create a lock for a specific repository path"""
        if repo_path not in self.repo_locks:
            self.repo_locks[repo_path] = threading.Lock()
        return self.repo_locks[repo_path]
        
    def _get_repo(self, repo_path: str):
        """Get Git repository, creating it if it doesn't exist"""
        if not self.git:
            raise ValueError("GitPython is not installed")
            
        if not os.path.exists(repo_path):
            os.makedirs(repo_path, exist_ok=True)
            
        try:
            repo = self.git.Repo(repo_path)
            return repo
        except self.git.InvalidGitRepositoryError:
            # Initialize new repository
            repo = self.git.Repo.init(repo_path)
            
            # Configure basic settings
            with repo.config_writer() as config:
                config.set_value("user", "name", self.default_author_name)
                config.set_value("user", "email", self.default_author_email)
                
            # Create initial commit
            if not os.path.exists(os.path.join(repo_path, '.gitignore')):
                with open(os.path.join(repo_path, '.gitignore'), 'w') as f:
                iles = [item.a_path for item in repo.index.diff("HEAD")]
                self.git_service.add_files(self.repo_path, ["README.md"])
                self.git_service.commit_changes(self.repo_path, "Initial repository setup")
        # Get untracked files
        # Large content threshold for chunking
        self.large_content_threshold = large_content_threshold
            "clean": not (staged_files or unstaged_files or untracked_files),
        # Index directory for document metadata
        self.index_dir = self.base_path / ".index"
        self.index_dir.mkdir(exist_ok=True)s,
            "untracked_files": untracked_files
        # Add vector search capability
        self.vector_search_enabled = False
        self.vector_model = None: str, file_path: Optional[str] = None, target: Optional[str] = None) -> str:
        try:et diff of changes"""
            import numpy as nprepo_path)
            from sentence_transformers import SentenceTransformer
            self.np = np  # Store numpy as instance variable
            self.vector_model = SentenceTransformer('all-MiniLM-L6-v2')
            self.vector_search_enabled = True_path)
            self.vector_index_path = self.base_path / ".vectors"
            self.vector_index_path.mkdir(exist_ok=True)
        except ImportError:
            logger.warning("Vector search dependencies not installed. Semantic search disabled.")

        # Add locks for shared resourcesfiles: List[str]) -> str:
        self.git_lock = Lock()mit"""
        self.index_lock = Lock()(repo_path):
        self.file_locks = {}_repo(repo_path)
            repo.git.add(files)
    def _get_file_lock(self, doc_id):essfully"
        """Get a lock for a specific file to prevent concurrent modifications"""
        if doc_id not in self.file_locks:tr, message: str, 
            self.file_locks[doc_id] = Lock()str] = None,
        return self.file_locks[doc_id]tional[str] = None) -> str:
        """Commit staged changes"""
    def create_document(self, ck(repo_path):
                       title: str, epo_path)
                       content: str,  or self.default_username
                       document_type: DocumentType,ault_email
                       metadata: Dict[str, Any] = None, 
                       tags: List[str] = None,g:
                       source_url: Optional[str] = None,name)
                       storage_type: str = "local") -> Dict[str, Any]:
        """Create a new document with Git versioning"""
        # Generate unique IDdex.commit(message)
        doc_id = f"doc_{int(time.time())}_{uuid.uuid4().hex[:8]}"
        now = int(time.time())
        reset_changes(self, repo_path: str) -> str:
        # Set up metadatahanges"""
        metadata = metadata or {}repo_path):
        tags = tags or []get_repo(repo_path)
            repo.git.reset()
        # Set up document directory based on type
        doc_dir = self.base_path / document_type.value
        doc_path = doc_dir / f"{doc_id}.md"ount: int = 10, file_path: Optional[str] = None) -> Dict[str, Any]:
        """Get log of commits"""
        # Prepare document content with frontmatter
        frontmatter = f"---\ntitle: {title}\ncreated_at: {now}\nupdated_at: {now}\nid: {doc_id}\ndocument_type: {document_type.value}\n"
            commits = list(repo.iter_commits(paths=file_path, max_count=max_count))
        # Add tags to frontmatter
        if tags:its = list(repo.iter_commits(max_count=max_count))
            frontmatter += f"tags: {', '.join(tags)}\n"
            commit in commits:
        # Add source URL if provided
        if source_url:: commit.hexsha,
            frontmatter += f"source_url: {source_url}\n"t.author.email}>",
                "date": commit.committed_datetime.strftime("%Y-%m-%d %H:%M:%S %z"),
        # Add any metadata to frontmatter.strip()
        for key, value in metadata.items():
            if isinstance(value, (str, int, float, bool)):
                frontmatter += f"{key}: {value}\n"
        create_branch(self, repo_path: str, branch_name: str, base_branch: Optional[str] = None) -> str:
        frontmatter += "---\n\n""
        full_content = frontmatter + content
            repo = self._get_repo(repo_path)
        # Write document to file
        doc_path.write_text(full_content, encoding='utf-8')
            repo.git.checkout('-b', branch_name)
        # Create document indexanch '{branch_name}'"
        self._update_index(doc_id, {
            "id": doc_id,elf, repo_path: str, branch_name: str, create: bool = False) -> str:
            "title": title,ting branch or create a new one"""
            "document_type": document_type.value,
            "created_at": now,epo(repo_path)
            "updated_at": now,
            "tags": tags,_name not in repo.refs:
            "metadata": metadata,kout('-b', branch_name)
            "size_bytes": len(full_content.encode('utf-8')),
            "source_url": source_url,f"Branch '{branch_name}' already exists")
            "path": str(doc_path.relative_to(self.base_path))
        })      repo.git.checkout(branch_name)
            return f"Checked out branch '{branch_name}'"
        # Add to git
        rel_path = doc_path.relative_to(self.base_path), auth_token: Optional[str] = None) -> str:
        with self.git_lock:sitory"""
            self.git_service.add_files(self.repo_path, [str(rel_path)])
            self.git_service.commit_changes(self.repo_path, f"Created document: {title}")
                if repo_url.startswith("https://"):
        # Update knowledge graph with document references", f"https://x-access-token:{auth_token}@")
        self._update_memory_graph(doc_id, title, document_type.value, tags, metadata, source_url)
            else:
        # Generate embeddings for semantic searchath)
        self.generate_embeddings(doc_id, content)_path}'"
        except Exception as e:
        return self.get_document(doc_id) clone repository: {str(e)}")
    
    def update_document(self, _path: str, file_path: str) -> str:
                       doc_id: str, pository"""
                       title: Optional[str] = None,
                       content: Optional[str] = None,
                       metadata: Optional[Dict[str, Any]] = None,
                       tags: Optional[List[str]] = None,
                       commit_message: str = "Updated document",
                       expected_version: Optional[str] = None) -> Dict[str, Any]:
        """Update an existing document with version control"""
        # Get file lock for this document to remove file: {str(e)}")
        with self._get_file_lock(doc_id):
            # Get current documento_path: str, file_path: str, version: str) -> str:
            doc_info = self._get_document_index(doc_id)version"""
            if not doc_info:o(repo_path)
                raise ValueError(f"Document with ID {doc_id} not found")
            blob = repo.git.show(f"{version}:{file_path}")
            # Add version check to prevent conflicts
            if expected_version:
                current_version = Noneto get file content at version {version}: {str(e)}")
                try:
                    log = self.git_service.get_log(tterns: List[str]) -> str:
                        self.repo_path,pository"""
                        max_count=1,o_path):
                        file_path=doc_info["path"]
                    )
                    if log.get("commits"):lfs", "install"])
                        current_version = log["commits"][0]["hash"]
                        .git.execute(["git", "lfs", "track", pattern])
                    if current_version and current_version != expected_version:
                        raise ValueError("Document has been modified since you loaded it. Please refresh and try again.")
                except Exception as e:
                    logger.warning(f"Version check failed: {e}")r(e)}")
                    
            # Get document path_path: str, file_groups: List[List[str]], message_template: str) -> List[str]:
            doc_path = self.base_path / doc_info["path"]e"""
            if not doc_path.exists():_path):
                raise ValueError(f"Document file not found at {doc_path}")
            commit_hashes = []
            # Read current contentnumerate(file_groups):
            current_content = doc_path.read_text(encoding='utf-8')
                commit = repo.index.commit(f"{message_template} (batch {i+1}/{len(file_groups)})")
            # Extract frontmatter and contentexsha)
            frontmatter_match = re.match(r"---(.*?)---\n\n", current_content, re.DOTALL)
            if not frontmatter_match:
                raise ValueError("Invalid document format: missing frontmatter") None, all_remotes: bool = False) -> str:
            ull changes from a remote repository"""
            frontmatter = frontmatter_match.group(1)
            existing_content = current_content[frontmatter_match.end():]
            try:
            # Update frontmatter
            frontmatter_dict = {}.git.fetch(all=True)
            for line in frontmatter.strip().split('\n'):
                if ': ' in line:o.git.pull(remote, branch)
                    key, value = line.split(': ', 1)
                    frontmatter_dict[key] = value
                raise ValueError(f"Failed to pull changes: {str(e)}")
            # Update fields
            now = int(time.time()): str, tag_name: str, message: str = None, commit: str = "HEAD") -> str:
            frontmatter_dict["updated_at"] = str(now)
             self._get_repo_lock(repo_path):
            if title:lf._get_repo(repo_path)
                frontmatter_dict["title"] = title
                if message:
            # Update tags if providedag_name, ref=commit, message=message)
            if tags is not None:
                if tags:.create_tag(tag_name, ref=commit)
                    frontmatter_dict["tags"] = ', '.join(tags)
                else:ception as e:
                    if "tags" in frontmatter_dict:e tag: {str(e)}")
                        del frontmatter_dict["tags"]
            _tags(self, repo_path: str) -> List[Dict[str, str]]:
            # Update metadatahe repository"""
            if metadata:_repo(repo_path)
                for key, value in metadata.items():
                    if isinstance(value, (str, int, float, bool)):
                        frontmatter_dict[key] = str(value)
                tags.append({
            # Build new frontmattere,
            new_frontmatter = "---\n"mit.hexsha,
            for key, value in frontmatter_dict.items():me.strftime("%Y-%m-%d %H:%M:%S %z")
                new_frontmatter += f"{key}: {value}\n"
            new_frontmatter += "---\n\n"
            pt Exception as e:
            # Update content if provided list tags: {str(e)}")
            final_content = content if content is not None else existing_content
            full_content = new_frontmatter + final_content
            ptimize the Git repository"""
            # Write updated document (protected by file lock)
            doc_path.write_text(full_content, encoding='utf-8')
            repo.git.gc('--aggressive', '--prune=now')
            # Update document index with index lockly"
            doc_info_update = {
                "updated_at": now,led to optimize repository: {str(e)}")
                "size_bytes": len(full_content.encode('utf-8')),
            }gure_auth(self, repo_path: str, username: str, password: str) -> str:
            onfigure authentication for repository operations"""
            if title:me or not password:
                doc_info_update["title"] = titleord required for HTTPS authentication")
                Storing passwords in git config is not secure
            if tags is not None:dential store or credential manager
                doc_info_update["tags"] = tags
                 = self._get_repo(repo_path)
            if metadata:nfig_writer() as config:
                doc_info_update["metadata"] = {**doc_info.get("metadata", {}), **metadata}
                config.set_value("user", "password", password)
            with self.index_lock:n configured successfully"
                self._update_index(doc_id, doc_info_update)
            ster_webhook(self, repo_path: str, webhook: Dict[str, Any]) -> str:
            # Git operations with git lockts"""
            with self.git_lock:holder implementation. In a real-world scenario, you would need to handle webhooks
                rel_path = doc_path.relative_to(self.base_path)
                self.git_service.add_files(self.repo_path, [str(rel_path)])
                self.git_service.commit_changes(self.repo_path, commit_message)
                _file.write(f"#!/bin/sh\ncurl -X POST {webhook['url']} -d @- <<'EOF'\n$(git log -1 --pretty=format:'%H')\nEOF\n")
            # Update memory graph)
            self._update_memory_graph(cessfully"
                doc_id,
                title or doc_info.get("title"),tr, file_path: str, version: str) -> bool:
                doc_info.get("document_type"),n"""
                tags if tags is not None else doc_info.get("tags", []),
                {**doc_info.get("metadata", {}), **(metadata or {})},
                doc_info.get("source_url")t_at_version(repo_path, file_path, version)
            )
            # Write that content to the current file
            # Generate embeddings for semantic searchpath)
            if content is not None:name(full_path), exist_ok=True)
                self.generate_embeddings(doc_id, content)
            with open(full_path, 'w', encoding='utf-8') as f:
        return self.get_document(doc_id)
            
    def get_document(self, doc_id: str) -> Dict[str, Any]:
        """Get document metadata and preview"""h])
        doc_info = self._get_document_index(doc_id)ed file to version {version}")
        if not doc_info:
            return None  # This line was missing
            logger.error(f"Error restoring file version: {e}", exc_info=True)
        # Get document path
        doc_path = self.base_path / doc_info["path"]
        if not doc_path.exists():
            return Noneeld(..., description="Path to the Git repository")
        
        # Read content for preview):
        content = doc_path.read_text(encoding='utf-8') files to add")
        age: str = Field(..., description="Commit message")
        # Strip frontmatter for previewNone, description="Author name")
        content_without_frontmatter = re.sub(r"^---.*?---\n\n", "", content, flags=re.DOTALL)
        preview = content_without_frontmatter[:500] + ("..." if len(content_without_frontmatter) > 500 else "")
        tDiffRequest(GitRepoPath):
        # Get version count from git(None, description="Path to the file to diff")
        version_count = 1 = Field(None, description="Target to diff against")
        try:
            log = self.git_service.get_log(
                self.repo_path,description="Maximum number of commits to return")
                max_count=100, Field(None, description="Path to the file to get log for")
                file_path=doc_info["path"]
            )chRequest(GitRepoPath):
            version_count = len(log.get("commits", []))the branch to create")
        except Exception:[str] = Field(None, description="Base branch to create the new branch from")
            pass
        tCheckoutRequest(GitRepoPath):
        return { str = Field(..., description="Name of the branch to checkout")
            "id": doc_id,False, description="Create the branch if it doesn't exist")
            "title": doc_info.get("title", "Untitled"),
            "document_type": doc_info.get("document_type", DocumentType.GENERIC.value),
            "created_at": doc_info.get("created_at", 0),epository to clone")
            "updated_at": doc_info.get("updated_at", 0),one the repository to")
            "tags": doc_info.get("tags", []),escription="Authentication token for private repositories")
            "metadata": doc_info.get("metadata", {}),
            "content_preview": preview,:
            "size_bytes": doc_info.get("size_bytes", 0), file to remove")
            "version_count": version_count,
            "content_available": True,h):
            "source_url": doc_info.get("source_url") the file")
        }on: str = Field(..., description="Git version to get the file content from")
    
    def get_document_content(self, doc_id: str, version: Optional[str] = None) -> Dict[str, Any]:
        """Get full document content, optionally from a specific version"""s to track with LFS")
        doc_info = self._get_document_index(doc_id)
        if not doc_info:est(GitRepoPath):
            return Noneist[str]] = Field(..., description="List of file groups to commit in batches")
        age_template: str = Field(..., description="Template for commit messages")
        # Get document path
        doc_path = self.base_path / doc_info["path"]
        te: str = Field("origin", description="Remote to pull from")
        content = ""[str] = Field(None, description="Branch to pull")
        if version:ol = Field(False, description="Fetch from all remotes")
            # Get content from specific git version
            try:est(GitRepoPath):
                content = self.git_service.get_file_content_at_version(
                    self.repo_path,None, description="Tag message")
                    doc_info["path"],iption="Commit to tag")
                    version
                )onse(BaseModel):
            except Exception:] = Field(..., description="List of tags")
                return None
        else:ook(BaseModel):
            # Get current contenttion="Webhook URL")
            if not doc_path.exists():scription="List of events to trigger the webhook")
                return None Field(None, description="Webhook secret")
            content = doc_path.read_text(encoding='utf-8')
         APIRouter()
        # Strip frontmatter for content
        content_without_frontmatter = re.sub(r"^---.*?---\n\n", "", content, flags=re.DOTALL)
        
        return {tatus", response_model=Dict[str, Any], summary="Get repository status", description="Get the status of a Git repository, including the current branch, staged changes, and unstaged changes.")
            "id": doc_id,est: GitRepoPath = Body(...)):
            "title": doc_info.get("title", "Untitled"),
            "content": content_without_frontmatter,
            "version": version_status(request.repo_path)
        }t ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    def get_document_versions(self, doc_id: str, max_versions: int = 10) -> List[Dict[str, Any]]:
        """Get version history for a document"""ail=f"Failed to get status: {str(e)}")
        doc_info = self._get_document_index(doc_id)
        if not doc_info:sponse_model=str, summary="Get diff of changes", description="Get the difference between working directory and HEAD or a specified target.")
            return []quest: GitDiffRequest = Body(...)):
            he difference between working directory and HEAD or a specified target."""
        try:
            log = self.git_service.get_log(.repo_path, request.file_path, request.target)
                self.repo_path,
                max_count=max_versions,=400, detail=str(e))
                file_path=doc_info["path"]
            ) HTTPException(status_code=500, detail=f"Failed to get diff: {str(e)}")
            
            versions = []onse_model=str, summary="Stage files for commit", description="Stage files for commit.")
            for commit in log.get("commits", []):ody(...)):
                versions.append({
                    "version_hash": commit["hash"],
                    "commit_message": commit["message"],request.files)
                    "author": commit["author"],
                    "timestamp": int(time.mktime(time.strptime(commit["date"], "%Y-%m-%d %H:%M:%S %z")))
                })on as e:
            e HTTPException(status_code=500, detail=f"Failed to add files: {str(e)}")
            return versions
        except Exception as e:se_model=str, summary="Commit changes", description="Commit staged changes with a commit message. Optionally, specify the author name and email.")
            logger.error(f"Error getting versions: {e}", exc_info=True)
            return []changes with a commit message. Optionally, specify the author name and email."""
    try:
    def delete_document(self, doc_id: str) -> bool:epo_path, request.message, request.author_name, request.author_email)
        """Delete a document"""
        doc_info = self._get_document_index(doc_id)=str(e))
        if not doc_info:e:
            return Falseion(status_code=500, detail=f"Failed to commit changes: {str(e)}")
        
        # Get document pathonse_model=str, summary="Reset staged changes", description="Reset all staged changes.")
        doc_path = self.base_path / doc_info["path"]...)):
        if not doc_path.exists():""
            return False
        return git_service.reset_changes(request.repo_path)
        try:alueError as e:
            # Remove fileon(status_code=400, detail=str(e))
            doc_path.unlink()
            e HTTPException(status_code=500, detail=f"Failed to reset changes: {str(e)}")
            # Remove from git
            rel_path = doc_path.relative_to(self.base_path)ummary="Get commit log", description="Get the commit log of the repository")
            self.git_service.remove_file(self.repo_path, str(rel_path))
            self.git_service.commit_changes(self.repo_path, f"Deleted document: {doc_info.get('title', doc_id)}")
            
            # Remove from indexlog(request.repo_path, request.max_count, request.file_path)
            self._remove_index(doc_id)
            e HTTPException(status_code=400, detail=str(e))
            # Remove from memory graph
            self._remove_from_memory_graph(doc_id)l=f"Failed to get log: {str(e)}")
            
            return True response_model=str, summary="Create branch", description="Create a new branch from a base branch.")
        except Exception as e:t: GitBranchRequest = Body(...)):
            logger.error(f"Error deleting document: {e}", exc_info=True)
            return False
        return git_service.create_branch(request.repo_path, request.branch_name, request.base_branch)
    def search_documents(self, 
                        query: str, ode=400, detail=str(e))
                        doc_type: Optional[str] = None, 
                        tags: Optional[List[str]] = None,led to create branch: {str(e)}")
                        limit: int = 10) -> List[Dict[str, Any]]:
        """Search documents by query, type, and tags"""Checkout branch", description="Checkout an existing branch or create a new one.")
        results = []ranch(request: GitCheckoutRequest = Body(...)):
        heckout an existing branch or create a new one."""
        # Get all document IDs
        index_files = list(self.index_dir.glob("*.json"))ath, request.branch_name, request.create)
        pt ValueError as e:
        for index_file in index_files:e=400, detail=str(e))
            try:tion as e:
                doc_info = json.loads(index_file.read_text(encoding='utf-8'))ch: {str(e)}")
                
                # Apply filters_model=str, summary="Clone repository", description="Clone a Git repository.")
                if doc_type and doc_info.get("document_type") != doc_type:
                    continuey."""
                
                if tags:ce.clone_repo(request.repo_url, request.local_path, request.auth_token)
                    doc_tags = set(doc_info.get("tags", []))
                    if not all(tag in doc_tags for tag in tags):
                        continue
                TPException(status_code=500, detail=f"Failed to clone repository: {str(e)}")
                # Check if query matches
                if query:esponse_model=str, summary="Remove file", description="Remove a file from the repository.")
                    query_lower = query.lower()uest = Body(...)):
                    title = doc_info.get("title", "").lower()
                    doc_content = ""
                    ervice.remove_file(request.repo_path, request.file_path)
                    # Only load content if necessary for search
                    if query and not (query_lower in title):
                        try:
                            doc_path = self.base_path / doc_info["path"]ile: {str(e)}")
                            content = doc_path.read_text(encoding='utf-8')
                            # Strip frontmatterr, summary="Get file content", description="Get the content of a file at a specific Git version.")
                            doc_content = re.sub(r"^---.*?---\n\n", "", content, flags=re.DOTALL).lower()
                        except Exception:cific Git version."""
                            pass
                    ervice.get_file_content(request.repo_path, request.file_path, request.version)
                    if not (query_lower in title or query_lower in doc_content):
                        # No match in title or contentr(e))
                        continue
                TPException(status_code=500, detail=f"Failed to get file content: {str(e)}")
                # Format result
                results.append({del=str, summary="Configure Git LFS", description="Configure Git LFS for the repository.")
                    "id": doc_info.get("id"),t = Body(...)):
                    "title": doc_info.get("title", "Untitled"),
                    "document_type": doc_info.get("document_type", DocumentType.GENERIC.value),
                    "created_at": doc_info.get("created_at", 0),est.file_patterns)
                    "updated_at": doc_info.get("updated_at", 0),
                    "tags": doc_info.get("tags", []),tr(e))
                    "metadata": doc_info.get("metadata", {}),
                    "size_bytes": doc_info.get("size_bytes", 0),set up Git LFS: {str(e)}")
                    "source_url": doc_info.get("source_url")
                })ch-commit", response_model=List[str], summary="Batch commit", description="Commit files in batches for better performance.")
                commit(request: GitBatchCommitRequest = Body(...)):
                # Limit resultsfor better performance."""
                if len(results) >= limit:
                    breake.batch_commit(request.repo_path, request.file_groups, request.message_template)
                    r as e:
            except Exception as e:_code=400, detail=str(e))
                logger.error(f"Error processing document index {index_file}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to batch commit: {str(e)}")
        return results
    ter.post("/pull", response_model=str, summary="Pull changes", description="Pull changes from a remote repository.")
    def _update_index(self, doc_id: str, doc_info: Dict[str, Any]) -> None:
        """Update the document index"""itory."""
        with self.index_lock:
            index_path = self.index_dir / f"{doc_id}.json" request.remote, request.branch, request.all_remotes)
            alueError as e:
            if index_path.exists():code=400, detail=str(e))
                current_info = json.loads(index_path.read_text(encoding='utf-8'))
                current_info.update(doc_info)detail=f"Failed to pull changes: {str(e)}")
                index_path.write_text(json.dumps(current_info, indent=2), encoding='utf-8')
            else:g", response_model=str, summary="Create tag", description="Create a new Git tag.")
                # Create new indexagRequest = Body(...)):
                index_path.write_text(json.dumps(doc_info, indent=2), encoding='utf-8')
    try:
    def _get_document_index(self, doc_id: str) -> Optional[Dict[str, Any]]:request.message, request.commit)
        """Get document index by ID"""
        index_path = self.index_dir / f"{doc_id}.json"r(e))
        if not index_path.exists():
            return Nonetion(status_code=500, detail=f"Failed to create tag: {str(e)}")
        
        try:("/tags", response_model=GitTagsResponse, summary="List tags", description="List all tags in the repository.")
            return json.loads(index_path.read_text(encoding='utf-8'))
        except Exception:he repository."""
            return None
        tags = git_service.list_tags(request.repo_path)
    def _remove_index(self, doc_id: str) -> None:
        """Remove document index"""
        index_path = self.index_dir / f"{doc_id}.json"r(e))
        if index_path.exists():
            index_path.unlink()tus_code=500, detail=f"Failed to list tags: {str(e)}")
    
    def _update_memory_graph(self,_model=str, summary="Optimize repository", description="Optimize the Git repository.")
                           doc_id: str, Path = Body(...)):
                           title: str, 
                           doc_type: str,
                           tags: List[str], uest.repo_path)
                           metadata: Dict[str, Any],
                           source_url: Optional[str] = None) -> None:
        """Update the knowledge graph with document references"""
        try:e HTTPException(status_code=500, detail=f"Failed to optimize repository: {str(e)}")
            # Create or update document entity
            doc_entity_name = f"document:{doc_id}"
            it__(self, data_dir: str = None):
            # Create observations
            observations = [ories
                f"Title: {title}",dir or os.path.join(os.getcwd(), 'data', 'scraped'))
                f"Type: {doc_type}",_dir / "records"
            ]data_dir.mkdir(parents=True, exist_ok=True)
            .records_dir.mkdir(parents=True, exist_ok=True)
            if tags:ing configuration
                observations.append(f"Tags: {', '.join(tags)}")
                _delay = config.scraper_max_delay
            if source_url:config.user_agent
                observations.append(f"Source URL: {source_url}")
            bot parser cache
            # Add metadata as observations
            for key, value in metadata.items():
                if isinstance(value, (str, int, float, bool)):
                    observations.append(f"{key}: {value}")
            ._browser_lock = asyncio.Lock()
            # Add document entity to graph
            self.memory_service.create_entities([{
                "name": doc_entity_name, per-domain delay settings
                "entity_type": "document",
                "observations": observationsstr, min_delay: float, max_delay: float):
            }])specific rate limits for a domain"""
            .domain_delays[domain] = (min_delay, max_delay)
            # Create relations for tags
            relations = []elf, url: str) -> Tuple[float, float]:
            for tag in tags: delay for specific domain"""
                # Create tag entity if needed
                tag_entity_name = f"tag:{tag}"
                self.memory_service.create_entities([{
                    "name": tag_entity_name,
                    "entity_type": "tag",ain]
                    "observations": [f"Document tag: {tag}"]
                }]) = '.'.join(domain.split('.')[-2:])  # Get example.com from subdomain
                base_domain}" in self.domain_delays:
                # Create relationlays[f"*.{base_domain}"]
                relations.append({
                    "from": doc_entity_name,ay)
                    "to": tag_entity_name,
                    "relation_type": "tagged_with"
                })rowser instance, creating one if needed"""
            c with self._browser_lock:
            # Create relations for source if provided
            if source_url: = await async_playwright().start()
                source_entity_name = f"source:{source_url.replace('://', '_').replace('/', '_')}"
                self.memory_service.create_entities([{
                    "name": source_entity_name,
                    "entity_type": "source",robotparser.RobotFileParser:
                    "observations": [f"URL: {source_url}"]"
                }])= urlparse(url)
                 f"{parsed_url.scheme}://{parsed_url.netloc}"
                relations.append({
                    "from": doc_entity_name,
                    "to": source_entity_name,)
                    "relation_type": "sourced_from"
                })
                rp.set_url(robots_url)
            # Add relations to graph
            if relations:f we got a crawl delay
                self.memory_service.create_relations(relations)
                if delay:
        except Exception as e:elay = max(self.min_delay, delay)
            logger.error(f"Error updating memory graph: {e}", exc_info=True)
            except Exception as e:
    def _remove_from_memory_graph(self, doc_id: str) -> None:m {domain}: {e}")
        """Remove document from knowledge graph"""llback
        try:    rp = robotparser.RobotFileParser()
            doc_entity_name = f"document:{doc_id}"
            self.memory_service.delete_entities([doc_entity_name])
        except Exception as e:
            logger.error(f"Error removing document from memory graph: {e}", exc_info=True)
                         wait_for_selector: Optional[str] = None,
    def generate_embeddings(self, doc_id: str, content: str) -> None:
        """Generate and store embeddings for a document""", Any]:
        if not self.vector_search_enabled:logic"""
            returnt in range(retry_count):
            try:
        try:    # Initialize browser with Playwright
            # Generate embeddingself.get_browser()
            embedding = self.vector_model.encode(content[:10000])  # Limit size for efficiency
                    user_agent=self.user_agent,
            # Save embedding={'width': 1920, 'height': 1080}
            embedding_path = self.vector_index_path / f"{doc_id}.npy"
            self.np.save(embedding_path, embedding)locking analytics/ads
        except Exception as e:route("**/{analytics,ads,tracking}*.{js,png,jpg}", lambda route: route.abort())
            logger.error(f"Error generating embeddings for document {doc_id}: {str(e)}")
                page = await context.new_page()
    def semantic_search(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search documents semantically using vector similarity"""
        if not self.vector_search_enabled:ext)
            raise ValueError("Vector search not enabled")
                try:
        try:        # Check robots.txt
            # Generate query embeddingf.get_robot_parser(url)
            query_embedding = self.vector_model.encode(query)gent, url):
                        raise ValueError(f"Access to {url} is disallowed by robots.txt")
            # Compare with all document embeddings
            results = []vigate to page with appropriate waiting
            for embedding_file in self.vector_index_path.glob("*.npy"): timeout=30000)
                doc_id = embedding_file.stem
                doc_embedding = self.np.load(embedding_file)
                            await page.wait_for_selector(wait_for_selector, timeout=wait_for_timeout or 30000)
                # Calculate cosine similarity:
                similarity = self.np.dot(query_embedding, doc_embedding) / (not found: {e}")
                    self.np.linalg.norm(query_embedding) * self.np.linalg.norm(doc_embedding)
                )   
                    # Implement rate limiting
                results.append((doc_id, float(similarity))).min_delay, self.max_delay))
                    
            # Sort by similarity and get top results
            results.sort(key=lambda x: x[1], reverse=True)
            top_results = results[:limit]html, "lxml")
                    title = soup.title.string.strip() if soup.title and soup.title.string else "Untitled Page"
            # Get document information for top results
            return [self.get_document(doc_id) for doc_id, _ in top_results]
        except Exception as e: self._extract_metadata(soup, url)
            logger.error(f"Error during semantic search: {str(e)}")
            return [] Extract tables if present
                    tables = self._extract_tables(soup)
    def convert_document_format(self, doc_id: str, target_format: str) -> bytes:
        """Convert document to different formats (PDF, DOCX, etc.)"""
        # Get document contentontent = html_to_markdown(html, url, title)
        doc_content = self.get_document_content(doc_id)
        if not doc_content:record
            raise ValueError(f"Document {doc_id} not found")uid4().hex[:8]}"
                    record_path = self.records_dir / f"{record_id}.json"
        content = doc_content["content"]
        title = doc_content["title"]
                        "title": title,
        if target_format.lower() == "pdf":n_content,
            try:        "metadata": metadata,
                # Check if markdown is installed first,
                if markdown is None:rue
                    raise ImportError("markdown package is required for PDF conversion")
                    if tables:
                # Try importing weasyprint only when needed
                try:
                    import weasyprint
                except ImportError:iles.open(record_path, "w", encoding="utf-8") as f:
                    raise ImportError("weasyprint package is required for PDF conversion")
                    
                html_content = f"<h1>{title}</h1>{markdown.markdown(content)}"
                pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
                return pdf_bytes(f"Error scraping {url}: {e}")
            except ImportError as e:
                logger.error(f"{e}")
                raise ValueError(str(e))
                        "content": "",
        elif target_format.lower() == "docx":
            try:        "scraped_at": int(time.time()),
                # Try importing docx only when needed
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("python-docx package is required for DOCX conversion")
                
                doc = Document() e:
                doc.add_heading(title, 0)t {attempt + 1} failed for {url}: {e}")
                doc.add_paragraph(content)= retry_count:
                
                # Save to bytes
                buffer = io.BytesIO(): "",
                doc.save(buffer)
                buffer.seek(0)
                return buffer.read()                "scraped_at": int(time.time()),
            except ImportError as e:           "success": False,
                logger.error(f"{e}")
                raise ValueError(str(e))                    }
        
        else:as_documents: bool = False) -> List[Dict[str, Any]]:
            raise ValueError(f"Unsupported format: {target_format}")
rape_url(url) for url in urls]
    def get_document_diff(self, doc_id: str, from_version: str, to_version: str = "HEAD") -> Dict[str, Any]:
        """Get differences between document versions"""
        doc_info = self._get_document_index(doc_id)f crawl_website(self, 
        if not doc_info:
            raise ValueError(f"Document with ID {doc_id} not found")
        ursion_depth: int = 2,
        try:ed_domains: Optional[List[str]] = None,
            # Use git diff to get differenceserification_pass: bool = False) -> Dict[str, Any]:
            diff = self.git_service.get_file_diff( starting from a URL"""
                self.repo_path,se the start URL to get the base domain
        start_parsed = urlparse(start_url)
        start_domain = start_parsed.netloc
        
        # Set up allowed domains
        if allowed_domains is None:
            allowed_domains = [start_domain]
        
        # Track visited URLs and frontier
        visited = set()
        frontier = [(start_url, 0)]  # (url, depth)
        results = []
        
        # Set up the robot parser
        robot_parser = self.get_robot_parser(start_url)
        
        # Start crawling
        browser = await self.get_browser()
        
        while frontier and len(visited) < max_pages:
            url, depth = frontier.pop(0)
            if url in visited:
                continue
            
            # Check if URL is allowed
            parsed = urlparse(url)
            if parsed.netloc not in allowed_domains:
                continue
            
            # Check robots.txt
            if not robot_parser.can_fetch(self.user_agent, url):
                continue
            
            logger.info(f"Crawling: {url} (depth {depth})")
            visited.add(url)
            
            # Scrape the page
            result = await self.scrape_url(url)
            results.append(result)
            
            # Extract links for next level if not at max depth
            if depth < recursion_depth and result["success"]:
                # Extract links from HTML
                try:
                    links = self._extract_links(url, result["content"])
                    # Filter links and add to frontier
                    for link in links:
                        link_parsed = urlparse(link)
                        # Skip if already visited or queued
                        if link in visited or any(link == f[0] for f in frontier):
                            continue
                        
                        # Skip if not in allowed domains
                        if link_parsed.netloc not in allowed_domains:
                            continue
                        
                        # Add to frontier
                        frontier.append((link, depth + 1))
                        
                        # Limit frontier size
                        if len(visited) + len(frontier) >= max_pages:
                            break
                except Exception as e:
                    logger.error(f"Error extracting links from {url}: {e}")
            
            # Implement rate limiting
            await asyncio.sleep(random.uniform(self.min_delay, self.max_delay))
        
        # Prepare response
        response = {
            "pages_crawled": len(visited),
            "start_url": start_url,
            "success_count": sum(1 for r in results if r.get("success", False)),
            "failed_count": sum(1 for r in results if not r.get("success", False)),
            "results": results
        }
        
        # Perform verification pass if requested
        if verification_pass and visited:
            logger.info("Starting verification pass...")
            verification_results = []
            
            # Sample a subset of pages for verification (10% or at most 10 pages)
            verification_sample_size = min(10, max(1, int(len(visited) * 0.1)))
            verification_urls = random.sample(list(visited), verification_sample_size)
            for sample_url in verification_urls:
                try:
                    logger.info(f"Verifying: {sample_url}")
                    verification = await self.scrape_url(sample_url)
                    
                    # Check if content matches original scrape
                    original_result = next((r for r in results if r["url"] == sample_url), None)
                    content_match = False
                    
                    if original_result and verification["success"]:
                        # Simple content length comparison as basic check
                        orig_len = len(original_result.get("content", ""))
                        verify_len = len(verification.get("content", ""))
                        content_match = abs(orig_len - verify_len) / max(orig_len, 1) < 0.1  # Within 10%
                    
                    verification_results.append({
                        "url": sample_url,
                        "verified": verification["success"],
                        "content_consistent": content_match
                    })
                    
                    # Apply rate limiting between verification requests
                    await asyncio.sleep(random.uniform(self.min_delay, self.max_delay))
                except Exception as e:
                    logger.error(f"Verification failed for {sample_url}: {e}")
                    verification_results.append({
                        "url": sample_url,
                        "verified": False,
                        "error": str(e)
                    })
            
            # Add verification results to response
            response["verification_results"] = verification_results
            response["verification_success_rate"] = sum(1 for v in verification_results if v["verified"]) / len(verification_results) if verification_results else 0
        
        return response
    
    async def search_and_scrape(self, query: str, max_results: int = 10) -> List[Dict[str, Any]]:
        """Search for content and scrape the results"""
        # Use multiple search engines with rotation for reliability
        search_engines = [
            f"https://duckduckgo.com/html/?q={query}",
            f"https://www.bing.com/search?q={query}"
        ]
        
        results = []
        for search_url in search_engines:
            try:
                response = requests.get(search_url, headers={"User-Agent": self.user_agent})
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, "lxml")
                    links = [a['href'] for a in soup.select("a.result__a")][:max_results]
                    results.extend(await self.scrape_urls(links))
                    if len(results) >= max_results:
                        break
            except Exception as e:
                logger.warning(f"Search failed on {search_url}: {str(e)}")
                # Try next search engine instead of failing
                continue
        
        # Return empty results only if all engines fail
        return results[:max_results]
    
    def _extract_metadata(self, soup: BeautifulSoup, url: str) -> Dict[str, Any]:
        """Extract metadata from HTML page"""
        metadata = {}
        
        # Extract Open Graph metadata
        for meta in soup.find_all("meta"):
            if meta.get("property") and meta.get("property").startswith("og:"):
                property_name = meta.get("property")[3:]
                content = meta.get("content")
                if property_name and content:
                    metadata[property_name] = content
        
        # Extract standard metadata
        for meta in soup.find_all("meta"):
            name = meta.get("name")
            content = meta.get("content")
            if name and content:
                metadata[name] = content
        
        # Extract publication date
        date_meta = soup.find("meta", {"property": "article:published_time"})
        if date_meta:
            metadata["publication_date"] = date_meta.get("content")
        
        # Get author information
        author_meta = soup.find("meta", {"name": "author"}) or soup.find("meta", {"property": "article:author"})
        if author_meta:
            metadata["author"] = author_meta.get("content")
                
        # Add URL domain as source
        parsed_url = urlparse(url)
        metadata["source_domain"] = parsed_url.netloc
        
        return metadata
    
    def _extract_tables(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:
        """Extract tables from HTML content"""
        tables = []
        
        for idx, table in enumerate(soup.find_all("table")):
            table_data = {
                "headers": [],
                "rows": []
            }
            
            # Try to extract headers
            headers = []
            header_row = table.find("thead")
            if header_row:
                header_cells = header_row.find_all(["th", "td"])
                headers = [cell.get_text(strip=True) for cell in header_cells]
            else:
                # If no thead, try to use the first row as headers
                first_row = table.find("tr")
                if first_row:
                    header_cells = first_row.find_all(["th", "td"])
                    headers = [cell.get_text(strip=True) for cell in header_cells]
                
            table_data["headers"] = headers
            
            # Extract rows
            rows = []
            data_rows = table.find_all("tr")
            # Skip the first row if we used it for headers
            start_idx = 1 if not header_row and headers and len(data_rows) > 0 else 0
                
            for row in data_rows[start_idx:]:
                cells = row.find_all(["td", "th"])
                row_data = [cell.get_text(strip=True) for cell in cells]
                
                # Only add non-empty rows
                if any(cell for cell in row_data):
                    rows.append(row_data)
            
            table_data["rows"] = rows
            
            # Only add tables with actual data
            if headers or rows:
                tables.append(table_data)
        
        return tables
    
    def _extract_links(self, base_url: str, content: str) -> List[str]:
        """Extract links from markdown content"""
        # Look for markdown links [text](url)
        link_pattern = r'\[.*?\]\((https?://[^)]+)\)'
        links = re.findall(link_pattern, content)
        
        # Add base URL for relative links
        absolute_links = []
        for link in links:
            if link.startswith("http"):
                absolute_links.append(link)
            else:
                absolute_links.append(urljoin(base_url, link))
        
        return absolute_links
    
    async def get_or_scrape_url(self, url: str, max_cache_age: int = 86400) -> Dict[str, Any]:
        """Get from cache or scrape if not available/expired"""
        cache_key = self._get_cache_key(url)
        cache_path = self.cache_dir / f"{cache_key}.json"
        
        # Check if cache exists and is valid
        if cache_path.exists():
            try:
                cache_data = json.loads(cache_path.read_text(encoding="utf-8"))
                cached_at = cache_data.get("scraped_at", 0)
                # Check if cache is still valid
                if time.time() - cached_at < max_cache_age:
                    return cache_data
            except:
                pass
        
        # Cache miss or expired, perform scrape
        result = await self.scrape_url(url)
        
        # Save to cache if successful
        if result["success"]:
            self.cache_dir.mkdir(exist_ok=True)
            with open(cache_path, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
        
        return result
    
    def _get_cache_key(self, url: str) -> str:
        """Generate a cache key from a URL"""
        # Remove protocol and query parameters for consistent caching
        normalized_url = re.sub(r'^https?://', '', url)
        normalized_url = re.sub(r'\?.*$', '', normalized_url)
        # Remove trailing slashes
        normalized_url = normalized_url.rstrip('/')
        # Create a hash to avoid file system issues with long URLs
        url_hash = hashlib.md5(url.encode()).hexdigest()
        # Return a file-system friendly key
        return f"{url_hash}_{re.sub(r'[^a-zA-Z0-9_-]', '_', normalized_url)[:50]}"
    
    def _configure_proxy(self, context) -> None:
        """Configure proxy for playwright browser context"""
        if not hasattr(self, 'proxy_settings') or not self.proxy_settings:
            return
        
        # Add proxy authentication if provided
        auth = None
        if self.proxy_settings.get('username') and self.proxy_settings.get('password'):
            auth = {
                'username': self.proxy_settings['username'],
                'password': self.proxy_settings['password']
            }
        
        # Configure the proxy
        proxy = {
            'server': self.proxy_settings['server'],
            'bypass': self.proxy_settings.get('bypass', '')
        }
        if auth:
            proxy['username'] = auth['username']
            proxy['password'] = auth['password']
        
        # Apply proxy settings to the context
        context.route('**/*', lambda route: route.continue_(proxy=proxy))
    
    def extract_structured_data(self, html_content: str) -> Dict[str, Any]:
        """Extract structured data from HTML page"""
        try:
            soup = BeautifulSoup(html_content, 'lxml')
            structured_data = {}
            
            # Extract JSON-LD
            json_ld_data = []
            for script in soup.find_all('script', type='application/ld+json'):
                try:
                    data = json.loads(script.string)
                    if isinstance(data, dict):
                        json_ld_data.append(data)
                    elif isinstance(data, list):
                        json_ld_data.extend(data)
                except Exception as e:
                    logger.error(f"Error parsing JSON-LD: {e}")
            if json_ld_data:
                structured_data['json_ld'] = json_ld_data
            
            # Extract microdata
            microdata = {}
            for element in soup.find_all(itemscope=True):
                if element.has_attr('itemtype'):
                    item_type = element['itemtype']
                    item_props = {}
                    
                    # Extract properties
                    for prop in element.find_all(itemprop=True):
                        prop_name = prop['itemprop']
                        # Get property value based on tag type
                        if prop.name == 'meta':
                            prop_value = prop.get('content', '')
                        elif prop.name == 'img':
                            prop_value = prop.get('src', '')
                        elif prop.name == 'a':
                            prop_value = prop.get('href', '')
                        elif prop.name == 'time':
                            prop_value = prop.get('datetime', prop.get_text())
                        else:
                            prop_value = prop.get_text().strip()
                        
                        item_props[prop_name] = prop_value
                    
                    if item_type not in microdata:
                        microdata[item_type] = []
                    microdata[item_type].append(item_props)
            if microdata:
                structured_data['microdata'] = microdata
            
            # Extract OpenGraph metadata
            og_data = {}
            for meta in soup.find_all('meta', property=re.compile(r'^og:')):
                prop = meta.get('property', '').replace('og:', '')
                content = meta.get('content', '')
                if prop and content:
                    og_data[prop] = content
            if og_data:
                structured_data['opengraph'] = og_data
            
            # Extract Twitter card metadata
            twitter_data = {}
            for meta in soup.find_all('meta', attrs={'name': re.compile(r'^twitter:')}):
                prop = meta.get('name', '').replace('twitter:', '')
                content = meta.get('content', '')
                if prop and content:
                    twitter_data[prop] = content
            if twitter_data:
                structured_data['twitter_card'] = twitter_data
            
            return structured_data
        except Exception as e:
            logger.error(f"Error extracting structured data: {e}")
            return {}
    
    async def _handle_rate_limiting(self, response):
        """Handle rate limiting based on response codes"""
        if response.status == 429:  # Too Many Requests
            retry_after = response.headers.get('retry-after')
            wait_time = int(retry_after) if retry_after and retry_after.isdigit() else 60
            logger.info(f"Rate limited. Waiting for {wait_time} seconds")
            await asyncio.sleep(wait_time)
            return True
        
        return False
    
    async def capture_screenshot(self, url: str, full_page: bool = True) -> Dict[str, Any]:
        """Capture screenshot of a webpage"""
        try:
            browser = await self.get_browser()
            context = await browser.new_context(
                user_agent=self.user_agent,
                viewport={'width': 1920, 'height': 1080}
            )
            page = await context.new_page()
            await page.goto(url, wait_until="networkidle")
            
            # Capture screenshot
            screenshot_path = self.data_dir / "screenshots"
            screenshot_path.mkdir(exist_ok=True)
            filename = f"{hashlib.md5(url.encode()).hexdigest()}.png"
            file_path = screenshot_path / filename
            await page.screenshot(path=str(file_path), full_page=full_page)
            
            return {
                "url": url,
                "screenshot_path": str(file_path),
                "timestamp": int(time.time()),
                "success": True
            }
        except Exception as e:
            logger.error(f"Screenshot failed for {url}: {e}")
            return {
                "url": url,
                "error": str(e),
                "success": False
            }
    
    async def scrape_with_pagination(self, url: str, max_pages: int = 5) -> Dict[str, Any]:
        """Scrape a URL and follow pagination links"""
        all_content = ""
        current_url = url
        pages_scraped = 0
        
        while current_url and pages_scraped < max_pages:
            result = await self.scrape_url(current_url)
            if not result["success"]:
                break
            
            # Accumulate content
            all_content += result["content"] + "\n\n---\n\n"
            pages_scraped += 1
            
            # Find next page link
            next_url = self._find_next_page_link(result["content"], current_url)
            if not next_url or next_url == current_url:
                break
            
            current_url = next_url
            # Add delay between pages
            await asyncio.sleep(random.uniform(self.min_delay, self.max_delay))
        
        # Create combined result
        return {
            "url": url,
            "title": f"Paginated content ({pages_scraped} pages)",
            "content": all_content,
            "scraped_at": int(time.time()),
            "success": True,
            "pages_scraped": pages_scraped
        }
    
    def _find_next_page_link(self, content: str, current_url: str) -> Optional[str]:
        """Find pagination link in content"""
        soup = BeautifulSoup(content, "lxml")
        
        # Common patterns for next page links
        next_selectors = [
            '.pagination .next',
            '.pagination a[rel="next"]',
            'a.next',
            'a:contains("Next")',
            'a[aria-label="Next"]',
            '.pagination a:contains("›")',
            '.pagination a:contains(">")'
        ]
        
        for selector in next_selectors:
            try:
                next_link = soup.select_one(selector)
                if next_link and next_link.get('href'):
                    return urljoin(current_url, next_link['href'])
            except:
                continue
        
        return None
    
    def _extract_main_content(self, soup: BeautifulSoup) -> str:
        """Extract the main content from a page, ignoring navigation, ads, etc."""
        # Remove common non-content elements
        for element in soup.select('header, footer, nav, aside, .ads, .comments, .sidebar, script, style'):
            element.decompose()
        
        # Try to find the main content container
        main_content = None
        for selector in ['article', 'main', '.content', '.post', '#content', '[itemprop="articleBody"]']:
            content = soup.select_one(selector)
            if content:
                main_content = content
                break
        
        if not main_content:
            # Fall back to density-based content extraction if no container found
            paragraphs_by_parent = {}
            for p in soup.find_all('p'):
                parent = p.parent
                if parent not in paragraphs_by_parent:
                    paragraphs_by_parent[parent] = []
                paragraphs_by_parent[parent].append(p)
            
            max_text_length = 0
            main_content = soup.body or soup
            for parent, paragraphs in paragraphs_by_parent.items():
                text_length = sum(len(p.get_text()) for p in paragraphs)
                if text_length > max_text_length:
                    main_content = parent
                    max_text_length = text_length
        
        return main_content.get_text(separator="\n").strip()