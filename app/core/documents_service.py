import asyncio
import hashlib
import json
import osutil
import tempfile
from typing import List, Optional, Dict, Any, Union
import timeomimport Path
import randommport List, Dict, Any, Optional, Tuple
import gitging
from git import Repo
import reng import List, Optional, Dict, Any, Union, Set, Tuple
from fastapi import APIRouter, Body, HTTPException
import uuid.parse import urljoin, urlparseng
from pydantic import BaseModel, Fieldser
from urllib.parse import urljoin, urlparse
import urllib.robotparser as robotparser
import aiofilesport APIRouter, Body, Query, HTTPException, Path, UploadFile, File, Form, Response
from pathlib import PatheModel, Field not installed. Markdown to HTML conversion will be limited.")
from typing import List, Dict, Any, Optional, Set, Tuple
import loggings.documents import (mport FilesystemService
import threading,ory_service import MemoryService
from bs4 import BeautifulSouport GitService
from playwright.async_api import async_playwright, Playwright, Browser, PagesionResponse, DocumentContentResponse
import requestsponse, import get_config
from app.utils.markdown import html_to_markdown
from app.utils.config import get_config
)ogger = logging.getLogger(__name__)
logger = logging.getLogger(__name__)rt DocumentsService
from app.utils.config import get_config
class GitService:self, 
    def __init__(self):ath: str = None, 
        config = get_config()name__)None,
        self.default_username = config.default_git_username ~100KB for large documents
        self.default_email = config.default_git_email
        self.temp_auth_files = {}
        self.repo_locks = {}    
    tags=["Documents"],paths
    def _get_repo(self, repo_path: str) -> git.Repo:etcwd(), 'data', 'documents'))
        """Get git repository object"""not found"},parents=True, exist_ok=True)
        try: {"description": "Operation failed"}
            repo = Repo(repo_path)
            return repo
        except git.exc.InvalidGitRepositoryError:
            raise ValueError(f"Invalid Git repository at '{repo_path}'")
        except Exception as e:ervice()
            raise ValueError(f"Failed to get repository: {str(e)}")
@router.post(.git_service = GitService()
    def _get_repo_lock(self, repo_path: str) -> threading.Lock:
        """Get a lock for a specific repository to prevent concurrent modifications"""file_path)
        if repo_path not in self.repo_locks:
            self.repo_locks[repo_path] = threading.Lock()g. Supports manuscripts, documentation, and datasets."
        return self.repo_locks[repo_path]
async def create_document(request: CreateDocumentRequest):
    def get_status(self, repo_path: str) -> Dict[str, Any]:
        """Get the status of a Git repository"""
        repo = self._get_repo(repo_path)
        current_branch = repo.active_branch.nameatasets") as f:
        # Get staged filesnts (70,000+ words)naged by Tools Server Document Service.")
        staged_files = [item.a_path for item in repo.index.diff("HEAD")]
        # Get modified but unstaged filesdd_files(self.repo_path, ["README.md"])
        unstaged_files = [item.a_path for item in repo.index.diff(None)]
        # Get untracked filesservice.create_document(
        untracked_files = repo.untracked_files
        return {ent=request.content,tent_threshold
            "clean": not (staged_files or unstaged_files or untracked_files),
            metadata=request.metadata,
            tags=request.tags, ".index"
            source_url=request.source_url
        ): untracked_files
        return document
    except Exception as e:tor_search_enabled = False
        raise HTTPException(status_code=500, detail=f"Document creation failed: {str(e)}")t: Optional[str] = None) -> str:

@router.get(
    "/{doc_id}", SentenceTransformer
    response_model=DocumentResponse,path) variable
    summary="Get document by ID",lif file_path:_model = SentenceTransformer('all-MiniLM-L6-v2')
    description="Retrieve a document by its unique identifier"            return repo.git.diff('HEAD', file_path)
)
async def get_document(doc_id: str = Path(..., description="Document unique identifier")):target)(exist_ok=True)
    """Retrieve a document by its ID"""
    document = documents_service.get_document(doc_id)() search dependencies not installed. Semantic search disabled.")
    if not document:
        raise HTTPException(status_code=404, detail=f"Document {doc_id} not found") repo_path: str, files: List[str]) -> str:
    return document
et_repo_lock(repo_path):
@router.put(ath)
    "/{doc_id}",epo.git.add(files)
    response_model=DocumentResponse,uccessfully"
    summary="Update document",        """Get a lock for a specific file to prevent concurrent modifications"""
    description="Update an existing document with version control"
)ptional[str] = None,
async def update_document([str] = None) -> str:
    doc_id: str = Path(..., description="Document unique identifier"),
    request: UpdateDocumentRequest = Body(...)k(repo_path):
):
    """Update an existing document with version control"""            author_name = author_name or self.default_username
    try:il
        document = documents_service.update_document(
            doc_id=doc_id,
            title=request.title,ser", "name", author_name)
            content=request.content,ail", author_email), Any]:
            metadata=request.metadata,rsioning"""
            tags=request.tags,
            commit_message=request.commit_messageha}"
        )
        if not document:tr:
            raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
        return document
    except ValueError as e:_repo(repo_path)
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document update failed: {str(e)}")        doc_dir = self.base_path / document_type.value
t = 10, file_path: Optional[str] = None) -> Dict[str, Any]:
@router.get(
    "/{doc_id}/content",tmatter
    response_model=DocumentContentResponse,ncreated_at: {now}\nupdated_at: {now}\nid: {doc_id}\ndocument_type: {document_type.value}\n"
    summary="Get document content",epo.iter_commits(paths=file_path, max_count=max_count))
    description="Get the full content of a document"
)            commits = list(repo.iter_commits(max_count=max_count))
async def get_document_content(
    doc_id: str = Path(..., description="Document unique identifier"),
    version: Optional[str] = Query(None, description="Specific version to retrieve")
):": commit.hexsha,
    """Get the full content of a document, optionally from a specific version"""
    content = documents_service.get_document_content(doc_id, version)   "date": commit.committed_datetime.strftime("%Y-%m-%d %H:%M:%S %z"),
    if not content:
        raise HTTPException(status_code=404, detail=f"Document content not found")e in metadata.items():
    return contente, (str, int, float, bool)):
= f"{key}: {value}\n"
@router.get( str, branch_name: str, base_branch: Optional[str] = None) -> str:
    "/{doc_id}/versions",
    response_model=List[DocumentVersionResponse],
    summary="Get document versions",
    description="Get version history for a document" base_branch: file
)t.checkout(base_branch)ding='utf-8')
async def get_document_versions(            repo.git.checkout('-b', branch_name)
    doc_id: str = Path(..., description="Document unique identifier"),
    max_versions: int = Query(10, description="Maximum number of versions to return"), {
):, branch_name: str, create: bool = False) -> str:
    """Get version history for a document"""eate a new one"""
    versions = documents_service.get_document_versions(doc_id, max_versions)_lock(repo_path):alue,
    return versions

@router.delete(
    "/{doc_id}",                    repo.git.checkout('-b', branch_name)
    response_model=Dict[str, Any],
    summary="Delete document", already exists")
    description="Delete a document"(self.base_path))
)e)
async def delete_document(doc_id: str = Path(..., description="Document unique identifier")):hecked out branch '{branch_name}'"
    """Delete a document"""
    success = documents_service.delete_document(doc_id) auth_token: Optional[str] = None) -> str:
    if not success:t repository"""
        raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
    return {"success": True, "message": f"Document {doc_id} deleted"}th_token:ce.commit_changes(self.repo_path, f"Created document: {title}")
//"):
@router.get(/", f"https://x-access-token:{auth_token}@")
    "/search",                Repo.clone_from(repo_url, local_path)ment_type.value, tags, metadata, source_url)
    response_model=List[DocumentResponse],
    summary="Search documents",_url, local_path)
    description="Search documents by query, type, and tags"return f"Cloned repository to '{local_path}'"
)s e:
async def search_documents(sitory: {str(e)}")
    query: str = Query("", description="Search query"),
    doc_type: Optional[str] = Query(None, description="Document type filter"),str) -> str:
    tags: Optional[str] = Query(None, description="Comma-separated tags to filter by"), a file from the repository"""
    limit: int = Query(10, description="Maximum results to return")
):
    """Search documents by query, type, and tags"""a: Optional[Dict[str, Any]] = None,
    tag_list = tags.split(",") if tags else None
    results = documents_service.search_documents(query, doc_type, tag_list, limit)                repo.index.commit(f"Removed {file_path}")ument",
    return resultst"ict[str, Any]:
sion control"""
@router.get( remove file: {str(e)}")
    "/{doc_id}/convert",
    summary="Convert format",_content(self, repo_path: str, file_path: str, version: str) -> str:
    description="Convert document to different format (PDF, DOCX)",ific Git version"""
    responses={
        200: {und")
            "content": {"application/pdf": {}, "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {}},"{version}:{file_path}")
            "description": "Converted document"
        }        except Exception as e:n:
    }(e)}")
)
async def convert_document(tr, file_patterns: List[str]) -> str:
    doc_id: str = Path(..., description="Document unique identifier"),onfigure Git LFS for the repository"""
    format: str = Query(..., description="Target format (pdf, docx)")
):._get_repo(repo_path)path"]
    """Convert document to a different format (PDF, DOCX)"""
    try:
        # Get converted document bytes                for pattern in file_patterns:["commits"][0]["hash"]
        content_bytes = documents_service.convert_document_format(doc_id, format)
        tracking")= expected_version:
        # Set content type based on formatuccessfully"een modified since you loaded it. Please refresh and try again.")
        if format.lower() == "pdf":
            media_type = "application/pdf"raise ValueError(f"Failed to set up Git LFS: {str(e)}")
            filename = f"{doc_id}.pdf"
        elif format.lower() == "docx":le_groups: List[List[str]], message_template: str) -> List[str]:
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{doc_id}.docx"
        else:nd at {doc_path}")
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
        
        # Return response with proper content type                repo.git.add(file_group)ead_text(encoding='utf-8')
        return Response(
            content=content_bytes,
            media_type=media_type,--(.*?)---\n\n", current_content, re.DOTALL)
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )o_path: str, remote: str = "origin", branch: str = None, all_remotes: bool = False) -> str:
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))th):group(1)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Format conversion failed: {str(e)}")

@router.get(                    result = repo.git.fetch(all=True)
    "/{doc_id}/diff",
    response_model=Dict[str, Any],branch)
    summary="Compare document versions",': ', 1)
    description="Get differences between document versions"value
)raise ValueError(f"Failed to pull changes: {str(e)}")
async def get_document_diff(
    doc_id: str = Path(..., description="Document unique identifier"),r, message: str = None, commit: str = "HEAD") -> str:
    from_version: str = Query(..., description="Source version hash"),ew Git tag"""ted_at"] = str(now)
    to_version: str = Query("HEAD", description="Target version hash (default: current)")
):repo(repo_path)
    """Compare two versions of a document to see changes"""title"] = title
    try:
        diff = documents_service.get_document_diff(doc_id, from_version, to_version)                    repo.create_tag(tag_name, ref=commit, message=message)
        return diff
    except ValueError as e:g(tag_name, ref=commit)
        raise HTTPException(status_code=404, detail=str(e))ame}'"'.join(tags)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to compare versions: {str(e)}")raise ValueError(f"Failed to create tag: {str(e)}")
 frontmatter_dict["tags"]
class GitService:
    def __init__(self):ags in the repository"""
        config = get_config()
        self.default_username = config.default_git_username:
        self.default_email = config.default_git_emailvalue, (str, int, float, bool)):
        self.temp_auth_files = {}
        self.repo_locks = {}                tags.append({

    def _get_repo(self, repo_path: str) -> git.Repo:ha,
        """Get git repository object"""ommitted_datetime.strftime("%Y-%m-%d %H:%M:%S %z")
        try:    })w_frontmatter += f"{key}: {value}\n"
            repo = Repo(repo_path)gstter += "---\n\n"
            return repo
        except git.exc.InvalidGitRepositoryError:f"Failed to list tags: {str(e)}")
            raise ValueError(f"Invalid Git repository at '{repo_path}'")f content is not None else existing_content
        except Exception as e:tr:content
            raise ValueError(f"Failed to get repository: {str(e)}")
f._get_repo(repo_path)otected by file lock)
    def _get_repo_lock(self, repo_path: str) -> threading.Lock:ite_text(full_content, encoding='utf-8')
        """Get a lock for a specific repository to prevent concurrent modifications"""ressive', '--prune=now')
        if repo_path not in self.repo_locks:
            self.repo_locks[repo_path] = threading.Lock()        except Exception as e:{
        return self.repo_locks[repo_path]epository: {str(e)}")
ntent.encode('utf-8')),
    def get_status(self, repo_path: str) -> Dict[str, Any]:str, username: str, password: str) -> str:
        """Get the status of a Git repository"""onfigure authentication for repository operations"""
        repo = self._get_repo(repo_path)
        current_branch = repo.active_branch.namequired for HTTPS authentication")
        # Get staged filesrds in git config is not secure
        staged_files = [item.a_path for item in repo.index.diff("HEAD")]
        # Get modified but unstaged files        with self._get_repo_lock(repo_path):gs
        unstaged_files = [item.a_path for item in repo.index.diff(None)]
        # Get untracked files
        untracked_files = repo.untracked_files "name", username)get("metadata", {}), **metadata}
        return {
            "clean": not (staged_files or unstaged_files or untracked_files),
            "current_branch": current_branch,
            "staged_files": staged_files,r, webhook: Dict[str, Any]) -> str:
            "unstaged_files": unstaged_files,"""
            "untracked_files": untracked_filestion. In a real-world scenario, you would need to handle webhooks
        }path)
post-commit")
    def get_diff(self, repo_path: str, file_path: Optional[str] = None, target: Optional[str] = None) -> str:ath, commit_message)
        """Get diff of changes"""            hook_file.write(f"#!/bin/sh\ncurl -X POST {webhook['url']} -d @- <<'EOF'\n$(git log -1 --pretty=format:'%H')\nEOF\n")
        repo = self._get_repo(repo_path)
        if file_path and target:"
            return repo.git.diff(target, file_path)
        elif file_path:str, file_path: str, version: str) -> str:
            return repo.git.diff('HEAD', file_path)
        elif target:s if tags is not None else doc_info.get("tags", []),
            return repo.git.diff(target)
        else:"{version}:{file_path}")
            return repo.git.diff()
        except Exception as e:
    def add_files(self, repo_path: str, files: List[str]) -> str:)
        """Stage files for commit"""
        with self._get_repo_lock(repo_path):ore_file_version(self, repo_path: str, file_path: str, version: str) -> bool:    self.generate_embeddings(doc_id, content)
            repo = self._get_repo(repo_path)
            repo.git.add(files)
            return "Files staged successfully"version
o_path, file_path, version) -> Dict[str, Any]:
    def commit_changes(self, repo_path: str, message: str, 
                      author_name: Optional[str] = None,
                      author_email: Optional[str] = None) -> str:join(repo_path, file_path):
        """Commit staged changes"""(full_path), exist_ok=True)s line was missing
        with self._get_repo_lock(repo_path):
            repo = self._get_repo(repo_path)
            author_name = author_name or self.default_username(content)        doc_path = self.base_path / doc_info["path"]
            author_email = author_email or self.default_email
            # Set author for this commit
            with repo.config_writer() as config:es(repo_path, [file_path])        
                config.set_value("user", "name", author_name)            self.commit_changes(repo_path, f"Restored file to version {version}")
                config.set_value("user", "email", author_email)
            # Commit changes
            commit = repo.index.commit(message)            logger.error(f"Error restoring file version: {e}", exc_info=True)
            return f"Committed changes with hash {commit.hexsha}".DOTALL)
ontent_without_frontmatter) > 500 else "")
    def reset_changes(self, repo_path: str) -> str:
        """Reset staged changes"""")
        with self._get_repo_lock(repo_path):
            repo = self._get_repo(repo_path)class GitCommitRequest(GitRepoPath):        try:
            repo.git.reset()description="List of files to add")e.get_log(
            return "All staged changes reset"

    def get_log(self, repo_path: str, max_count: int = 10, file_path: Optional[str] = None) -> Dict[str, Any]:    author_email: Optional[str] = Field(None, description="Author email")                file_path=doc_info["path"]
        """Get log of commits"""
        repo = self._get_repo(repo_path)
        if file_path:
            commits = list(repo.iter_commits(paths=file_path, max_count=max_count))    target: Optional[str] = Field(None, description="Target to diff against")            pass
        else:
            commits = list(repo.iter_commits(max_count=max_count))
        log_data = []
        for commit in commits:    file_path: Optional[str] = Field(None, description="Path to the file to get log for")            "title": doc_info.get("title", "Untitled"),
            log_data.append({ENERIC.value),
                "hash": commit.hexsha,
                "author": f"{commit.author.name} <{commit.author.email}>",
                "date": commit.committed_datetime.strftime("%Y-%m-%d %H:%M:%S %z"),    base_branch: Optional[str] = Field(None, description="Base branch to create the new branch from")
                "message": commit.message.strip()t("metadata", {}),
            })
        return log_data
_count": version_count,
    def create_branch(self, repo_path: str, branch_name: str, base_branch: Optional[str] = None) -> str:
        """Create a new branch"""
        with self._get_repo_lock(repo_path):e")
            repo = self._get_repo(repo_path)    local_path: str = Field(..., description="Path to clone the repository to")    
            if base_branch:e, description="Authentication token for private repositories"), doc_id: str, version: Optional[str] = None) -> Dict[str, Any]:
                repo.git.checkout(base_branch)
            repo.git.checkout('-b', branch_name)oc_id)
            return f"Created branch '{branch_name}'"    file_path: str = Field(..., description="Path to the file to remove")

    def checkout_branch(self, repo_path: str, branch_name: str, create: bool = False) -> str:
        """Checkout an existing branch or create a new one"""    file_path: str = Field(..., description="Path to the file")        # Get document path
        with self._get_repo_lock(repo_path):="Git version to get the file content from")/ doc_info["path"]
            repo = self._get_repo(repo_path)
            if create:
                if branch_name not in repo.refs:    file_patterns: List[str] = Field(..., description="List of file patterns to track with LFS")
                    repo.git.checkout('-b', branch_name)cific git version
                else:
                    raise ValueError(f"Branch '{branch_name}' already exists")le groups to commit in batches")ntent_at_version(
            else:ssages")
                repo.git.checkout(branch_name)
            return f"Checked out branch '{branch_name}'":                    version
to pull from")
    def clone_repo(self, repo_url: str, local_path: str, auth_token: Optional[str] = None) -> str:")
        """Clone a Git repository""" all remotes")                return None
        try:
            if auth_token:
                if repo_url.startswith("https://"):
                    repo_url = repo_url.replace("https://", f"https://x-access-token:{auth_token}@")    message: Optional[str] = Field(None, description="Tag message")
                Repo.clone_from(repo_url, local_path)D", description="Commit to tag")            content = doc_path.read_text(encoding='utf-8')
            else:
                Repo.clone_from(repo_url, local_path)
            return f"Cloned repository to '{local_path}'"")ontmatter = re.sub(r"^---.*?---\n\n", "", content, flags=re.DOTALL)
        except Exception as e:        
            raise ValueError(f"Failed to clone repository: {str(e)}")eModel):
description="Webhook URL")
    def remove_file(self, repo_path: str, file_path: str) -> str:    events: List[str] = Field(..., description="List of events to trigger the webhook")ntitled"),
        """Remove a file from the repository"""
        with self._get_repo_lock(repo_path):
            repo = self._get_repo(repo_path)
            try:
                repo.index.remove([file_path])sions: int = 10) -> List[Dict[str, Any]]:
                repo.index.commit(f"Removed {file_path}")
                return f"Successfully removed {file_path} from Git"ary="Get repository status", description="Get the status of a Git repository, including the current branch, staged changes, and unstaged changes.")        doc_info = self._get_document_index(doc_id)
            except Exception as e:st: GitRepoPath = Body(...)):
                raise ValueError(f"Failed to remove file: {str(e)}")
    try:
    def get_file_content(self, repo_path: str, file_path: str, version: str) -> str:
        """Get the content of a file at a specific Git version"""
        repo = self._get_repo(repo_path)
        try:pt Exception as e:
            blob = repo.git.show(f"{version}:{file_path}")_info["path"]
            return blob
        except Exception as e: of changes", description="Get the difference between working directory and HEAD or a specified target.")            
            raise ValueError(f"Failed to get file content at version {version}: {str(e)}"): GitDiffRequest = Body(...)):
""
    def configure_lfs(self, repo_path: str, file_patterns: List[str]) -> str:    try:
        """Configure Git LFS for the repository"""h": commit["hash"],
        with self._get_repo_lock(repo_path):
            repo = self._get_repo(repo_path)s_code=400, detail=str(e))": commit["author"],
            try:pt Exception as e:ime(commit["date"], "%Y-%m-%d %H:%M:%S %z")))
                repo.git.execute(["git", "lfs", "install"])ff: {str(e)}")
                for pattern in file_patterns:
                    repo.git.execute(["git", "lfs", "track", pattern])es for commit", description="Stage files for commit.")            return versions
                repo.index.commit("Set up Git LFS tracking")t: GitCommitRequest = Body(...)):
                return "Git LFS configured successfully"
            except Exception as e:    try:
                raise ValueError(f"Failed to set up Git LFS: {str(e)}")

    def batch_commit(self, repo_path: str, file_groups: List[List[str]], message_template: str) -> List[str]:
        """Commit files in batches for better performance"""pt Exception as e:
        with self._get_repo_lock(repo_path):
            repo = self._get_repo(repo_path)
            commit_hashes = [] changes", description="Commit staged changes with a commit message. Optionally, specify the author name and email.")        
            for i, file_group in enumerate(file_groups):equest: GitCommitRequest = Body(...)):
                repo.git.add(file_group)d email."""
                commit = repo.index.commit(f"{message_template} (batch {i+1}/{len(file_groups)})")    try:
                commit_hashes.append(commit.hexsha)mail)    return False
            return commit_hashes
code=400, detail=str(e))
    def pull_changes(self, repo_path: str, remote: str = "origin", branch: str = None, all_remotes: bool = False) -> str:pt Exception as e:
        """Pull changes from a remote repository"""d to commit changes: {str(e)}")k()
        with self._get_repo_lock(repo_path):
            repo = self._get_repo(repo_path)taged changes", description="Reset all staged changes.")            # Remove from git
            try:quest: GitRepoPath = Body(...)):
                if all_remotes:
                    result = repo.git.fetch(all=True)    try:f.repo_path, f"Deleted document: {doc_info.get('title', doc_id)}")
                else:
                    result = repo.git.pull(remote, branch)
                return resulttail=str(e))dex(doc_id)
            except Exception as e:pt Exception as e:
                raise ValueError(f"Failed to pull changes: {str(e)}")mory graph

    def create_tag(self, repo_path: str, tag_name: str, message: str = None, commit: str = "HEAD") -> str:ummary="Get commit log", description="Get the commit log of the repository")            
        """Create a new Git tag""" GitLogRequest = Body(...)):
        with self._get_repo_lock(repo_path):
            repo = self._get_repo(repo_path)    try:t: {e}", exc_info=True)
            try:
                if message:
                    repo.create_tag(tag_name, ref=commit, message=message)il=str(e))lf, 
                else:pt Exception as e:
                    repo.create_tag(tag_name, ref=commit)tr] = None, 
                return f"Created tag '{tag_name}'"
            except Exception as e: branch", description="Create a new branch from a base branch.")                        limit: int = 10) -> List[Dict[str, Any]]:
                raise ValueError(f"Failed to create tag: {str(e)}")quest: GitBranchRequest = Body(...)):

    def convert_document_format(self, doc_id: str, target_format: str) -> bytes:
        """Convert document to different formats (PDF, DOCX, etc.)"""
        # Get document content
        doc_content = self.get_document_content(doc_id)
        if not doc_content:
            raise ValueError(f"Document {doc_id} not found")
        
        content = doc_content["content"]
        title = doc_content["title"]
        
        if target_format.lower() == "pdf":
            try:
                # Check if markdown is installed first
                if markdown is None:
                    raise ImportError("markdown package is required for PDF conversion")
                
                # Try importing weasyprint only when needed
                try:
                    import weasyprint
                except ImportError:
                    raise ImportError("weasyprint package is required for PDF conversion")
                
                html_content = f"<h1>{title}</h1>{markdown.markdown(content)}"
                pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
                return pdf_bytes
            except ImportError as e:
                logger.error(f"{e}")
                raise ValueError(str(e))
            
        elif target_format.lower() == "docx":
            try:
                # Try importing docx only when needed
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("python-docx package is required for DOCX conversion")
                
                doc = Document()
                doc.add_heading(title, 0)
                doc.add_paragraph(content)
                
                # Save to bytes
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer.read()
            except ImportError as e:
                logger.error(f"{e}")
                raise ValueError(str(e))
        
        else:
            raise ValueError(f"Unsupported format: {target_format}")

        raise HTTPException(status_code=500, detail=f"Failed to create branch: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.create_branch(request.repo_path, request.branch_name, request.base_branch)    try:    """Create a new branch from a base branch."""async def create_branch(request: GitBranchRequest = Body(...)):@router.post("/branch", response_model=str, summary="Create branch", description="Create a new branch from a base branch.")        raise HTTPException(status_code=500, detail=f"Failed to get log: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.get_log(request.repo_path, request.max_count, request.file_path)    try:    """Get the commit log of the repository."""async def get_log(request: GitLogRequest = Body(...)):@router.post("/log", response_model=List[Dict[str, Any]], summary="Get commit log", description="Get the commit log of the repository")        raise HTTPException(status_code=500, detail=f"Failed to reset changes: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.reset_changes(request.repo_path)    try:    """Reset all staged changes."""async def reset_changes(request: GitRepoPath = Body(...)):@router.post("/reset", response_model=str, summary="Reset staged changes", description="Reset all staged changes.")        raise HTTPException(status_code=500, detail=f"Failed to commit changes: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.commit_changes(request.repo_path, request.message, request.author_name, request.author_email)    try:    """Commit staged changes with a commit message. Optionally, specify the author name and email."""async def commit_changes(request: GitCommitRequest = Body(...)):@router.post("/commit", response_model=str, summary="Commit changes", description="Commit staged changes with a commit message. Optionally, specify the author name and email.")        raise HTTPException(status_code=500, detail=f"Failed to add files: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.add_files(request.repo_path, request.files)    try:    """Stage files for commit."""async def add_files(request: GitCommitRequest = Body(...)):@router.post("/add", response_model=str, summary="Stage files for commit", description="Stage files for commit.")        raise HTTPException(status_code=500, detail=f"Failed to get diff: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.get_diff(request.repo_path, request.file_path, request.target)    try:    """Get the difference between working directory and HEAD or a specified target."""async def get_diff(request: GitDiffRequest = Body(...)):@router.post("/diff", response_model=str, summary="Get diff of changes", description="Get the difference between working directory and HEAD or a specified target.")        raise HTTPException(status_code=500, detail=f"Failed to get status: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.get_status(request.repo_path)    try:    """Get the status of a Git repository."""async def get_status(request: GitRepoPath = Body(...)):@router.post("/status", response_model=Dict[str, Any], summary="Get repository status", description="Get the status of a Git repository, including the current branch, staged changes, and unstaged changes.")git_service = GitService()router = APIRouter()    secret: Optional[str] = Field(None, description="Webhook secret")    events: List[str] = Field(..., description="List of events to trigger the webhook")    url: str = Field(..., description="Webhook URL")class GitWebhook(BaseModel):    tags: List[Dict[str, str]] = Field(..., description="List of tags")class GitTagsResponse(BaseModel):    commit: str = Field("HEAD", description="Commit to tag")    message: Optional[str] = Field(None, description="Tag message")    tag_name: str = Field(..., description="Tag name")class GitTagRequest(GitRepoPath):    all_remotes: bool = Field(False, description="Fetch from all remotes")    branch: Optional[str] = Field(None, description="Branch to pull")    remote: str = Field("origin", description="Remote to pull from")class GitPullRequest(GitRepoPath):    message_template: str = Field(..., description="Template for commit messages")    file_groups: List[List[str]] = Field(..., description="List of file groups to commit in batches")class GitBatchCommitRequest(GitRepoPath):    file_patterns: List[str] = Field(..., description="List of file patterns to track with LFS")class GitLFSRequest(GitRepoPath):    version: str = Field(..., description="Git version to get the file content from")    file_path: str = Field(..., description="Path to the file")class GitFileContentRequest(GitRepoPath):    file_path: str = Field(..., description="Path to the file to remove")class GitRemoveFileRequest(GitRepoPath):    auth_token: Optional[str] = Field(None, description="Authentication token for private repositories")    local_path: str = Field(..., description="Path to clone the repository to")    repo_url: str = Field(..., description="URL of the repository to clone")class GitCloneRequest(GitRepoPath):    create: bool = Field(False, description="Create the branch if it doesn't exist")    branch_name: str = Field(..., description="Name of the branch to checkout")class GitCheckoutRequest(GitRepoPath):    base_branch: Optional[str] = Field(None, description="Base branch to create the new branch from")    branch_name: str = Field(..., description="Name of the branch to create")class GitBranchRequest(GitRepoPath):    file_path: Optional[str] = Field(None, description="Path to the file to get log for")    max_count: int = Field(10, description="Maximum number of commits to return")class GitLogRequest(GitRepoPath):    target: Optional[str] = Field(None, description="Target to diff against")    file_path: Optional[str] = Field(None, description="Path to the file to diff")class GitDiffRequest(GitRepoPath):    author_email: Optional[str] = Field(None, description="Author email")    author_name: Optional[str] = Field(None, description="Author name")    message: str = Field(..., description="Commit message")    files: List[str] = Field(..., description="List of files to add")class GitCommitRequest(GitRepoPath):    repo_path: str = Field(..., description="Path to the Git repository")class GitRepoPath(BaseModel):            return False            logger.error(f"Error restoring file version: {e}", exc_info=True)        except Exception as e:            return True            self.commit_changes(repo_path, f"Restored file to version {version}")            self.add_files(repo_path, [file_path])            # Add and commit the change                f.write(content)            with open(full_path, 'w', encoding='utf-8') as f:            os.makedirs(os.path.dirname(full_path), exist_ok=True)            full_path = os.path.join(repo_path, file_path)            # Write that content to the current file            content = self.get_file_content_at_version(repo_path, file_path, version)            # Get the file content at the specified version        try:        """Restore a file to a specific version"""    def restore_file_version(self, repo_path: str, file_path: str, version: str) -> bool:        return "Webhook registered successfully"        os.chmod(hook_path, 0o755)            hook_file.write(f"#!/bin/sh\ncurl -X POST {webhook['url']} -d @- <<'EOF'\n$(git log -1 --pretty=format:'%H')\nEOF\n")        with open(hook_path, "w") as hook_file:        hook_path = os.path.join(repo_path, ".git", "hooks", "post-commit")        # using Git hooks or a custom implementation.        # Note: This is a placeholder implementation. In a real-world scenario, you would need to handle webhooks        """Register a webhook for Git events"""    def register_webhook(self, repo_path: str, webhook: Dict[str, Any]) -> str:            return "Authentication configured successfully"                config.set_value("user", "password", password)                config.set_value("user", "name", username)            with repo.config_writer() as config:            repo = self._get_repo(repo_path)        with self._get_repo_lock(repo_path):        # Consider using git credential store or credential manager        # Note: Storing passwords in git config is not secure            raise ValueError("Username and password required for HTTPS authentication")        if not username or not password:        """Configure authentication for repository operations"""    def configure_auth(self, repo_path: str, username: str, password: str) -> str:            raise ValueError(f"Failed to optimize repository: {str(e)}")        except Exception as e:            return "Repository optimized successfully"            repo.git.gc('--aggressive', '--prune=now')        try:        repo = self._get_repo(repo_path)        """Optimize the Git repository"""    def optimize_repo(self, repo_path: str) -> str:            raise ValueError(f"Failed to list tags: {str(e)}")        except Exception as e:            return tags                })                    "date": tag.commit.committed_datetime.strftime("%Y-%m-%d %H:%M:%S %z")                    "commit": tag.commit.hexsha,                    "name": tag.name,                tags.append({            for tag in repo.tags:            tags = []        try:        repo = self._get_repo(repo_path)        """List all tags in the repository"""    def list_tags(self, repo_path: str) -> List[Dict[str, str]]:    try:
        return git_service.create_branch(request.repo_path, request.branch_name, request.base_branch)# Get all document IDs
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create branch: {str(e)}")

@router.post("/checkout", response_model=str, summary="Checkout branch", description="Checkout an existing branch or create a new one.")                
async def checkout_branch(request: GitCheckoutRequest = Body(...)):
    """Checkout an existing branch or create a new one."""") != doc_type:
    try:
        return git_service.checkout_branch(request.repo_path, request.branch_name, request.create)        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))s = set(doc_info.get("tags", []))
    except Exception as e:ags):
        raise HTTPException(status_code=500, detail=f"Failed to checkout branch: {str(e)}")ntinue

@router.post("/clone", response_model=str, summary="Clone repository", description="Clone a Git repository.")                # Check if query matches
async def clone_repo(request: GitCloneRequest = Body(...)):
    """Clone a Git repository."""
    try:tle", "").lower()
        return git_service.clone_repo(request.repo_url, request.local_path, request.auth_token)            doc_content = ""
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))load content if necessary for search
    except Exception as e::
        raise HTTPException(status_code=500, detail=f"Failed to clone repository: {str(e)}")y:

@router.post("/remove", response_model=str, summary="Remove file", description="Remove a file from the repository.")                            content = doc_path.read_text(encoding='utf-8')
async def remove_file(request: GitRemoveFileRequest = Body(...)):
    """Remove a file from the repository.""" content, flags=re.DOTALL).lower()
    try:
        return git_service.remove_file(request.repo_path, request.file_path)                    pass
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))(query_lower in title or query_lower in doc_content):
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to remove file: {str(e)}")ntinue

@router.post("/file-content", response_model=str, summary="Get file content", description="Get the content of a file at a specific Git version.")                # Format result
async def get_file_content(request: GitFileContentRequest = Body(...)):
    """Get the content of a file at a specific Git version."""
    try:e", "Untitled"),
        return git_service.get_file_content(request.repo_path, request.file_path, request.version)            "document_type": doc_info.get("document_type", DocumentType.GENERIC.value),
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))d_at": doc_info.get("updated_at", 0),
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get file content: {str(e)}")ata": doc_info.get("metadata", {}),

@router.post("/lfs", response_model=str, summary="Configure Git LFS", description="Configure Git LFS for the repository.")                    "source_url": doc_info.get("source_url")
async def configure_lfs(request: GitLFSRequest = Body(...)):
    """Configure Git LFS for the repository."""
    try:
        return git_service.configure_lfs(request.repo_path, request.file_patterns)        if len(results) >= limit:
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to set up Git LFS: {str(e)}")or(f"Error processing document index {index_file}: {e}", exc_info=True)

@router.post("/batch-commit", response_model=List[str], summary="Batch commit", description="Commit files in batches for better performance.")        return results
async def batch_commit(request: GitBatchCommitRequest = Body(...)):
    """Commit files in batches for better performance.""" Any]) -> None:
    try:
        return git_service.batch_commit(request.repo_path, request.file_groups, request.message_template)with self.index_lock:
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


















































































        raise HTTPException(status_code=500, detail=f"Failed to create branch: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.create_branch(request.repo_path, request.branch_name, request.base_branch)    try:    """Create a new branch from a base branch."""async def create_branch(request: GitBranchRequest = Body(...)):@router.post("/branch", response_model=str, summary="Create branch", description="Create a new branch from a base branch.")        raise HTTPException(status_code=500, detail=f"Failed to get log: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.get_log(request.repo_path, request.max_count, request.file_path)    try:    """Get the commit log of the repository."""async def get_log(request: GitLogRequest = Body(...)):@router.post("/log", response_model=List[Dict[str, Any]], summary="Get commit log", description="Get the commit log of the repository")        raise HTTPException(status_code=500, detail=f"Failed to reset changes: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.reset_changes(request.repo_path)    try:    """Reset all staged changes."""async def reset_changes(request: GitRepoPath = Body(...)):@router.post("/reset", response_model=str, summary="Reset staged changes", description="Reset all staged changes.")        raise HTTPException(status_code=500, detail=f"Failed to commit changes: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.commit_changes(request.repo_path, request.message, request.author_name, request.author_email)    try:    """Commit staged changes with a commit message. Optionally, specify the author name and email."""async def commit_changes(request: GitCommitRequest = Body(...)):@router.post("/commit", response_model=str, summary="Commit changes", description="Commit staged changes with a commit message. Optionally, specify the author name and email.")        raise HTTPException(status_code=500, detail=f"Failed to add files: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.add_files(request.repo_path, request.files)    try:    """Stage files for commit."""async def add_files(request: GitCommitRequest = Body(...)):@router.post("/add", response_model=str, summary="Stage files for commit", description="Stage files for commit.")        raise HTTPException(status_code=500, detail=f"Failed to get diff: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.get_diff(request.repo_path, request.file_path, request.target)    try:    """Get the difference between working directory and HEAD or a specified target."""async def get_diff(request: GitDiffRequest = Body(...)):@router.post("/diff", response_model=str, summary="Get diff of changes", description="Get the difference between working directory and HEAD or a specified target.")        raise HTTPException(status_code=500, detail=f"Failed to get status: {str(e)}")    except Exception as e:        raise HTTPException(status_code=400, detail=str(e))    except ValueError as e:        return git_service.get_status(request.repo_path)    try:    """Get the status of a Git repository."""async def get_status(request: GitRepoPath = Body(...)):@router.post("/status", response_model=Dict[str, Any], summary="Get repository status", description="Get the status of a Git repository, including the current branch, staged changes, and unstaged changes.")git_service = GitService()router = APIRouter()    secret: Optional[str] = Field(None, description="Webhook secret")    events: List[str] = Field(..., description="List of events to trigger the webhook")    url: str = Field(..., description="Webhook URL")class GitWebhook(BaseModel):    tags: List[Dict[str, str]] = Field(..., description="List of tags")class GitTagsResponse(BaseModel):    commit: str = Field("HEAD", description="Commit to tag")    message: Optional[str] = Field(None, description="Tag message")    tag_name: str = Field(..., description="Tag name")class GitTagRequest(GitRepoPath):    all_remotes: bool = Field(False, description="Fetch from all remotes")    branch: Optional[str] = Field(None, description="Branch to pull")    remote: str = Field("origin", description="Remote to pull from")class GitPullRequest(GitRepoPath):    message_template: str = Field(..., description="Template for commit messages")    file_groups: List[List[str]] = Field(..., description="List of file groups to commit in batches")class GitBatchCommitRequest(GitRepoPath):    file_patterns: List[str] = Field(..., description="List of file patterns to track with LFS")class GitLFSRequest(GitRepoPath):    version: str = Field(..., description="Git version to get the file content from")    file_path: str = Field(..., description="Path to the file")class GitFileContentRequest(GitRepoPath):    file_path: str = Field(..., description="Path to the file to remove")class GitRemoveFileRequest(GitRepoPath):    auth_token: Optional[str] = Field(None, description="Authentication token for private repositories")    local_path: str = Field(..., description="Path to clone the repository to")    repo_url: str = Field(..., description="URL of the repository to clone")class GitCloneRequest(GitRepoPath):    create: bool = Field(False, description="Create the branch if it doesn't exist")    branch_name: str = Field(..., description="Name of the branch to checkout")class GitCheckoutRequest(GitRepoPath):    base_branch: Optional[str] = Field(None, description="Base branch to create the new branch from")    branch_name: str = Field(..., description="Name of the branch to create")class GitBranchRequest(GitRepoPath):    file_path: Optional[str] = Field(None, description="Path to the file to get log for")    max_count: int = Field(10, description="Maximum number of commits to return")class GitLogRequest(GitRepoPath):    target: Optional[str] = Field(None, description="Target to diff against")    file_path: Optional[str] = Field(None, description="Path to the file to diff")class GitDiffRequest(GitRepoPath):    author_email: Optional[str] = Field(None, description="Author email")    author_name: Optional[str] = Field(None, description="Author name")    message: str = Field(..., description="Commit message")    files: List[str] = Field(..., description="List of files to add")class GitCommitRequest(GitRepoPath):    repo_path: str = Field(..., description="Path to the Git repository")class GitRepoPath(BaseModel):            return False            logger.error(f"Error restoring file version: {e}", exc_info=True)        except Exception as e:            return True            self.commit_changes(repo_path, f"Restored file to version {version}")            self.add_files(repo_path, [file_path])            # Add and commit the change                f.write(content)            with open(full_path, 'w', encoding='utf-8') as f:            os.makedirs(os.path.dirname(full_path), exist_ok=True)            full_path = os.path.join(repo_path, file_path)            # Write that content to the current file            content = self.get_file_content_at_version(repo_path, file_path, version)            # Get the file content at the specified version        try:        """Restore a file to a specific version"""    def restore_file_version(self, repo_path: str, file_path: str, version: str) -> bool:        return "Webhook registered successfully"        os.chmod(hook_path, 0o755)            hook_file.write(f"#!/bin/sh\ncurl -X POST {webhook['url']} -d @- <<'EOF'\n$(git log -1 --pretty=format:'%H')\nEOF\n")        with open(hook_path, "w") as hook_file:        hook_path = os.path.join(repo_path, ".git", "hooks", "post-commit")        # using Git hooks or a custom implementation.        # Note: This is a placeholder implementation. In a real-world scenario, you would need to handle webhooks        """Register a webhook for Git events"""    def register_webhook(self, repo_path: str, webhook: Dict[str, Any]) -> str:            return "Authentication configured successfully"                config.set_value("user", "password", password)                config.set_value("user", "name", username)            with repo.config_writer() as config:            repo = self._get_repo(repo_path)        with self._get_repo_lock(repo_path):        # Consider using git credential store or credential manager        # Note: Storing passwords in git config is not secure            raise ValueError("Username and password required for HTTPS authentication")        if not username or not password:        """Configure authentication for repository operations"""    def configure_auth(self, repo_path: str, username: str, password: str) -> str:            raise ValueError(f"Failed to optimize repository: {str(e)}")        except Exception as e:            return "Repository optimized successfully"            repo.git.gc('--aggressive', '--prune=now')        try:        repo = self._get_repo(repo_path)        """Optimize the Git repository"""    def optimize_repo(self, repo_path: str) -> str:            raise ValueError(f"Failed to list tags: {str(e)}")        except Exception as e:            return tags                })                    "date": tag.commit.committed_datetime.strftime("%Y-%m-%d %H:%M:%S %z")                    "commit": tag.commit.hexsha,                    "name": tag.name,                tags.append({            for tag in repo.tags:            tags = []        try:        repo = self._get_repo(repo_path)        """List all tags in the repository"""    def list_tags(self, repo_path: str) -> List[Dict[str, str]]:    try:
        return git_service.create_branch(request.repo_path, request.branch_name, request.base_branch)# Get all document IDs
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create branch: {str(e)}")

@router.post("/checkout", response_model=str, summary="Checkout branch", description="Checkout an existing branch or create a new one.")                
async def checkout_branch(request: GitCheckoutRequest = Body(...)):
    """Checkout an existing branch or create a new one."""") != doc_type:
    try:
        return git_service.checkout_branch(request.repo_path, request.branch_name, request.create)        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))s = set(doc_info.get("tags", []))
    except Exception as e:ags):
        raise HTTPException(status_code=500, detail=f"Failed to checkout branch: {str(e)}")ntinue

@router.post("/clone", response_model=str, summary="Clone repository", description="Clone a Git repository.")                # Check if query matches
async def clone_repo(request: GitCloneRequest = Body(...)):
    """Clone a Git repository."""
    try:tle", "").lower()
        return git_service.clone_repo(request.repo_url, request.local_path, request.auth_token)            doc_content = ""
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))load content if necessary for search
    except Exception as e::
        raise HTTPException(status_code=500, detail=f"Failed to clone repository: {str(e)}")y:

@router.post("/remove", response_model=str, summary="Remove file", description="Remove a file from the repository.")                            content = doc_path.read_text(encoding='utf-8')
async def remove_file(request: GitRemoveFileRequest = Body(...)):
    """Remove a file from the repository.""" content, flags=re.DOTALL).lower()
    try:
        return git_service.remove_file(request.repo_path, request.file_path)                    pass
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))(query_lower in title or query_lower in doc_content):
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to remove file: {str(e)}")ntinue

@router.post("/file-content", response_model=str, summary="Get file content", description="Get the content of a file at a specific Git version.")                # Format result
async def get_file_content(request: GitFileContentRequest = Body(...)):
    """Get the content of a file at a specific Git version."""
    try:e", "Untitled"),
        return git_service.get_file_content(request.repo_path, request.file_path, request.version)            "document_type": doc_info.get("document_type", DocumentType.GENERIC.value),
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))d_at": doc_info.get("updated_at", 0),
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get file content: {str(e)}")ata": doc_info.get("metadata", {}),

@router.post("/lfs", response_model=str, summary="Configure Git LFS", description="Configure Git LFS for the repository.")                    "source_url": doc_info.get("source_url")
async def configure_lfs(request: GitLFSRequest = Body(...)):
    """Configure Git LFS for the repository."""
    try:
        return git_service.configure_lfs(request.repo_path, request.file_patterns)        if len(results) >= limit:
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to set up Git LFS: {str(e)}")or(f"Error processing document index {index_file}: {e}", exc_info=True)

@router.post("/batch-commit", response_model=List[str], summary="Batch commit", description="Commit files in batches for better performance.")        return results
async def batch_commit(request: GitBatchCommitRequest = Body(...)):
    """Commit files in batches for better performance.""" Any]) -> None:
    try:
        return git_service.batch_commit(request.repo_path, request.file_groups, request.message_template)with self.index_lock:
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to batch commit: {str(e)}")fo = json.loads(index_path.read_text(encoding='utf-8'))

@router.post("/pull", response_model=str, summary="Pull changes", description="Pull changes from a remote repository.")                index_path.write_text(json.dumps(current_info, indent=2), encoding='utf-8')
async def pull_changes(request: GitPullRequest = Body(...)):
    """Pull changes from a remote repository."""
    try:e_text(json.dumps(doc_info, indent=2), encoding='utf-8')
        return git_service.pull_changes(request.repo_path, request.remote, request.branch, request.all_remotes)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))ex by ID"""
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to pull changes: {str(e)}")exists():

@router.post("/tag", response_model=str, summary="Create tag", description="Create a new Git tag.")        
async def create_tag(request: GitTagRequest = Body(...)):
    """Create a new Git tag."""oding='utf-8'))
    try:
        return git_service.create_tag(request.repo_path, request.tag_name, request.message, request.commit)    return None
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))oc_id: str) -> None:
    except Exception as e:index"""
        raise HTTPException(status_code=500, detail=f"Failed to create tag: {str(e)}")
ts():
@router.post("/tags", response_model=GitTagsResponse, summary="List tags", description="List all tags in the repository.")
async def list_tags(request: GitRepoPath = Body(...)):    
    """List all tags in the repository."""
    try:
        tags = git_service.list_tags(request.repo_path) 
        return {"tags": tags}                   doc_type: str,
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))metadata: Dict[str, Any],
    except Exception as e:) -> None:
        raise HTTPException(status_code=500, detail=f"Failed to list tags: {str(e)}")ledge graph with document references"""

@router.post("/optimize", response_model=str, summary="Optimize repository", description="Optimize the Git repository.")        # Create or update document entity
async def optimize_repo(request: GitRepoPath = Body(...)):y_name = f"document:{doc_id}"
    """Optimize the Git repository."""
    try:ions
        return git_service.optimize_repo(request.repo_path)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail[f"Failed to optimize repository: {str(e)}"])
    
class ScraperService:.join(tags)}")
    def __init__(self, data_dir: str = None):
        config = get_config()
        # Set up data directories        observations.append(f"Source URL: {source_url}")
        self.data_dir = Path(data_dir or os.path.join(os.getcwd(), 'data', 'scraped'))
        self.records_dir = self.data_dir / "records"bservations
        self.data_dir.mkdir(parents=True, exist_ok=True)    for key, value in metadata.items():
        self.records_dir.mkdir(parents=True, exist_ok=True)nce(value, (str, int, float, bool)):
        # Rate limiting configurationions.append(f"{key}: {value}")
        self.min_delay = config.scraper_min_delay
        self.max_delay = config.scraper_max_delay    # Add document entity to graph
        self.user_agent = config.user_agententities([{
        
        # Robot parser cache                "entity_type": "document",
        self.robot_parsers = {}
        
        # Browser instance
        self._browser = None            # Create relations for tags
        self._browser_lock = asyncio.Lock()
        
        # Domain-specific rate limitingntity if needed
        self.domain_delays = {}  # Store per-domain delay settingse = f"tag:{tag}"
e.create_entities([{
    def set_domain_rate_limit(self, domain: str, min_delay: float, max_delay: float):ame,
        """Set specific rate limits for a domain"""
        self.domain_delays[domain] = (min_delay, max_delay)ag}"]

    def get_domain_delay(self, url: str) -> Tuple[float, float]:
        """Get rate limiting delay for specific domain"""
        parsed = urlparse(url)nd({
        domain = parsed.netloc
        # Check exact domain match                "to": tag_entity_name,
        if domain in self.domain_delays:d_with"
            return self.domain_delays[domain]
        # Check wildcard domain matches (*.example.com)
        base_domain = '.'.join(domain.split('.')[-2:])  # Get example.com from subdomainurce if provided
        if f"*.{base_domain}" in self.domain_delays:
            return self.domain_delays[f"*.{base_domain}"]eplace('/', '_')}"
        # Return default delayice.create_entities([{
        return (self.min_delay, self.max_delay)                "name": source_entity_name,
    
    async def get_browser(self) -> Browser:
        """Get a browser instance, creating one if needed"""
        async with self._browser_lock:
            if self._browser is None:        relations.append({
                playwright = await async_playwright().start()
                self._browser = await playwright.chromium.launch(headless=True)
            return self._browserrom"
    })
    def get_robot_parser(self, url: str) -> robotparser.RobotFileParser:
        """Get a robot parser for the given URL's domain"""ns to graph
        parsed_url = urlparse(url)
        domain = f"{parsed_url.scheme}://{parsed_url.netloc}"relations(relations)
        
        if domain not in self.robot_parsers:
            rp = robotparser.RobotFileParser()xc_info=True)
            robots_url = f"{domain}/robots.txt"
            try:
                rp.set_url(robots_url)
                rp.read()
                # Check if we got a crawl delayoc_id}"
                delay = rp.crawl_delay("*")tities([doc_entity_name])
                if delay:    except Exception as e:
                    self.min_delay = max(self.min_delay, delay) document from memory graph: {e}", exc_info=True)
                    self.max_delay = max(self.max_delay, delay)
            except Exception as e:None:
                logger.warning(f"Error reading robots.txt from {domain}: {e}")
                # Create a permissive parser as fallback
                rp = robotparser.RobotFileParser()
            self.robot_parsers[domain] = rp
        return self.robot_parsers[domain]
    
    async def scrape_url(self, url: str, tent[:10000])  # Limit size for efficiency
                         wait_for_selector: Optional[str] = None,
                         wait_for_timeout: Optional[int] = None,
                         retry_count: int = 3) -> Dict[str, Any]:ding_path = self.vector_index_path / f"{doc_id}.npy"
        """Scrape a single URL with retry logic"""
        for attempt in range(retry_count):
            try:er.error(f"Error generating embeddings for document {doc_id}: {str(e)}")
                # Initialize browser with Playwright
                browser = await self.get_browser()_search(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
                context = await browser.new_context(using vector similarity"""
                    user_agent=self.user_agent,
                    viewport={'width': 1920, 'height': 1080}e ValueError("Vector search not enabled")
                )
                # Enable JavaScript intercept for blocking analytics/ads
                await context.route("**/{analytics,ads,tracking}*.{js,png,jpg}", lambda route: route.abort())
                
                page = await context.new_page()
                e with all document embeddings
                # Add proxy if configured
                self._configure_proxy(context)
                tem
                try:g = self.np.load(embedding_file)
                    # Check robots.txt
                    robot_parser = self.get_robot_parser(url)
                    if not robot_parser.can_fetch(self.user_agent, url):
                        raise ValueError(f"Access to {url} is disallowed by robots.txt")oc_embedding)
                    
                    # Navigate to page with appropriate waiting
                    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                    if wait_for_selector:
                        try:get top results
                            await page.wait_for_selector(wait_for_selector, timeout=wait_for_timeout or 30000)verse=True)
                        except Exception as e:
                            logger.warning(f"Selector '{wait_for_selector}' not found: {e}")
                    await page.wait_for_load_state("networkidle", timeout=30000)cument information for top results
                    r doc_id, _ in top_results]
                    # Implement rate limiting
                    await asyncio.sleep(random.uniform(self.min_delay, self.max_delay))rror(f"Error during semantic search: {str(e)}")
                    
                    # Extract content
                    html = await page.content()ment_format(self, doc_id: str, target_format: str) -> bytes:
                    soup = BeautifulSoup(html, "lxml")formats (PDF, DOCX, etc.)"""
                    title = soup.title.string.strip() if soup.title and soup.title.string else "Untitled Page"
                    = self.get_document_content(doc_id)
                    # Extract metadata from page
                    metadata = self._extract_metadata(soup, url)
                    
                    # Extract tables if present"content"]
                    tables = self._extract_tables(soup)]
                    
                    # Convert to markdown
                    markdown_content = html_to_markdown(html, url, title)
                    
                    # Save record
                    record_id = f"{int(time.time())}_{uuid.uuid4().hex[:8]}"aise ImportError("markdown package is required for PDF conversion")
                    record_path = self.records_dir / f"{record_id}.json"
                    record = {hen needed
                        "url": url,
                        "title": title,
                        "content": markdown_content,
                        "metadata": metadata,n")
                        "scraped_at": int(time.time()),
                        "success": True<h1>{title}</h1>{markdown.markdown(content)}"
                    }.HTML(string=html_content).write_pdf()
                    if tables:
                        record["tables"] = tablesor:
                    print package is required for PDF conversion")
                    # Save the recordF conversion requires the weasyprint library")
                    async with aiofiles.open(record_path, "w", encoding="utf-8") as f:
                        await f.write(json.dumps(record, indent=2, ensure_ascii=False))docx":
                    
                    return recordlar
                except Exception as e:nt
                    logger.error(f"Error scraping {url}: {e}") Document()
                    return {heading(title, 0)
                        "url": url,)
                        "title": "",    
                        "content": "",
                        "metadata": {},
                        "scraped_at": int(time.time()),
                        "success": False,0)
                        "error": str(e))
                    }
                finally:ocx package is required for DOCX conversion")
                    await context.close()conversion requires the python-docx library")
            
            except Exception as e:
                logger.error(f"Attempt {attempt + 1} failed for {url}: {e}")ted format: {target_format}")
                if attempt + 1 == retry_count:
                    return {def get_document_diff(self, doc_id: str, from_version: str, to_version: str = "HEAD") -> Dict[str, Any]:
                        "url": url,
                        "title": "",c_id)
                        "content": "",
                        "metadata": {},h ID {doc_id} not found")
                        "scraped_at": int(time.time()),    
                        "success": False,
                        "error": str(e)s
                    }f(
    
    async def scrape_urls(self, urls: List[str], store_as_documents: bool = False) -> List[Dict[str, Any]]:
        """Scrape multiple URLs in parallel"""
        tasks = [self.scrape_url(url) for url in urls]
        return await asyncio.gather(*tasks)
    
    async def crawl_website(self,  the content changes (not frontmatter)
                            start_url: str,     # This would need additional processing to remove frontmatter diffs
                            max_pages: int = 50,
                            recursion_depth: int = 2,
                            allowed_domains: Optional[List[str]] = None,
                            verification_pass: bool = False) -> Dict[str, Any]:        "title": doc_info.get("title", "Untitled"),
        """Crawl a website starting from a URL"""ion,
        # Parse the start URL to get the base domainsion": to_version,
        start_parsed = urlparse(start_url)
        start_domain = start_parsed.netloc
        except Exception as e:
        # Set up allowed domainsgetting document diff: {e}", exc_info=True)
        if allowed_domains is None:(e)}")
            allowed_domains = [start_domain]                # Track visited URLs and frontier        visited = set()        frontier = [(start_url, 0)]  # (url, depth)        results = []                # Set up the robot parser        robot_parser = self.get_robot_parser(start_url)                # Start crawling        browser = await self.get_browser()                while frontier and len(visited) < max_pages:            url, depth = frontier.pop(0)            if url in visited:                continue                        # Check if URL is allowed            parsed = urlparse(url)            if parsed.netloc not in allowed_domains:                continue                        # Check robots.txt            if not robot_parser.can_fetch(self.user_agent, url):                continue                        logger.info(f"Crawling: {url} (depth {depth})")            visited.add(url)                        # Scrape the page            result = await self.scrape_url(url)            results.append(result)                        # Extract links for next level if not at max depth            if depth < recursion_depth and result["success"]:                # Extract links from HTML                try:                    links = self._extract_links(url, result["content"])                    # Filter links and add to frontier                    for link in links:                        link_parsed = urlparse(link)                        # Skip if already visited or queued                        if link in visited or any(link == f[0] for f in frontier):                            continue                                                # Skip if not in allowed domains                        if link_parsed.netloc not in allowed_domains:                            continue                                                # Add to frontier                        frontier.append((link, depth + 1))                                                # Limit frontier size                        if len(visited) + len(frontier) >= max_pages:                            break                except Exception as e:                    logger.error(f"Error extracting links from {url}: {e}")                        # Implement rate limiting            await asyncio.sleep(random.uniform(self.min_delay, self.max_delay))                # Prepare response        response = {            "pages_crawled": len(visited),            "start_url": start_url,            "success_count": sum(1 for r in results if r.get("success", False)),            "failed_count": sum(1 for r in results if not r.get("success", False)),            "results": results        }                # Perform verification pass if requested        if verification_pass and visited:            logger.info("Starting verification pass...")            verification_results = []                        # Sample a subset of pages for verification (10% or at most 10 pages)            verification_sample_size = min(10, max(1, int(len(visited) * 0.1)))            verification_urls = random.sample(list(visited), verification_sample_size)            for sample_url in verification_urls:                try:                    logger.info(f"Verifying: {sample_url}")                    verification = await self.scrape_url(sample_url)                                        # Check if content matches original scrape                    original_result = next((r for r in results if r["url"] == sample_url), None)                    content_match = False                                        if original_result and verification["success"]:                        # Simple content length comparison as basic check                        orig_len = len(original_result.get("content", ""))                        verify_len = len(verification.get("content", ""))                        content_match = abs(orig_len - verify_len) / max(orig_len, 1) < 0.1  # Within 10%                                        verification_results.append({                        "url": sample_url,                        "verified": verification["success"],                        "content_consistent": content_match                    })                                        # Apply rate limiting between verification requests                    await asyncio.sleep(random.uniform(self.min_delay, self.max_delay))                except Exception as e:                    logger.error(f"Verification failed for {sample_url}: {e}")                    verification_results.append({                        "url": sample_url,                        "verified": False,                        "error": str(e)                    })                        # Add verification results to response            response["verification_results"] = verification_results            response["verification_success_rate"] = sum(1 for v in verification_results if v["verified"]) / len(verification_results) if verification_results else 0                return response        async def search_and_scrape(self, query: str, max_results: int = 10) -> List[Dict[str, Any]]:        """Search for content and scrape the results"""        # Use multiple search engines with rotation for reliability        search_engines = [            f"https://duckduckgo.com/html/?q={query}",            f"https://www.bing.com/search?q={query}"        ]                results = []        for search_url in search_engines:            try:                response = requests.get(search_url, headers={"User-Agent": self.user_agent})                if response.status_code == 200:                    soup = BeautifulSoup(response.text, "lxml")                    links = [a['href'] for a in soup.select("a.result__a")][:max_results]                    results.extend(await self.scrape_urls(links))                    if len(results) >= max_results:                        break            except Exception as e:                logger.warning(f"Search failed on {search_url}: {str(e)}")                # Try next search engine instead of failing                continue                # Return empty results only if all engines fail        return results[:max_results]        def _extract_metadata(self, soup: BeautifulSoup, url: str) -> Dict[str, Any]:        """Extract metadata from HTML page"""        metadata = {}                # Extract Open Graph metadata        for meta in soup.find_all("meta"):            if meta.get("property") and meta.get("property").startswith("og:"):                property_name = meta.get("property")[3:]                content = meta.get("content")                if property_name and content:                    metadata[property_name] = content                # Extract standard metadata        for meta in soup.find_all("meta"):            name = meta.get("name")            content = meta.get("content")            if name and content:                metadata[name] = content                # Extract publication date        date_meta = soup.find("meta", {"property": "article:published_time"})        if date_meta:            metadata["publication_date"] = date_meta.get("content")                # Get author information        author_meta = soup.find("meta", {"name": "author"}) or soup.find("meta", {"property": "article:author"})        if author_meta:            metadata["author"] = author_meta.get("content")                        # Add URL domain as source        parsed_url = urlparse(url)        metadata["source_domain"] = parsed_url.netloc                return metadata        def _extract_tables(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:        """Extract tables from HTML content"""        tables = []                for idx, table in enumerate(soup.find_all("table")):            table_data = {                "headers": [],                "rows": []            }                        # Try to extract headers            headers = []            header_row = table.find("thead")            if header_row:                header_cells = header_row.find_all(["th", "td"])                headers = [cell.get_text(strip=True) for cell in header_cells]            else:                # If no thead, try to use the first row as headers                first_row = table.find("tr")                if first_row:                    header_cells = first_row.find_all(["th", "td"])                    headers = [cell.get_text(strip=True) for cell in header_cells]                            table_data["headers"] = headers                        # Extract rows            rows = []            data_rows = table.find_all("tr")            # Skip the first row if we used it for headers            start_idx = 1 if not header_row and headers and len(data_rows) > 0 else 0                            for row in data_rows[start_idx:]:                cells = row.find_all(["td", "th"])                row_data = [cell.get_text(strip=True) for cell in cells]                                # Only add non-empty rows                if any(cell for cell in row_data):                    rows.append(row_data)                        table_data["rows"] = rows                        # Only add tables with actual data            if headers or rows:                tables.append(table_data)                return tables        def _extract_links(self, base_url: str, content: str) -> List[str]:        """Extract links from markdown content"""        # Look for markdown links [text](url)        link_pattern = r'\[.*?\]\((https?://[^)]+)\)'        links = re.findall(link_pattern, content)                # Add base URL for relative links        absolute_links = []        for link in links:            if link.startswith("http"):                absolute_links.append(link)            else:                absolute_links.append(urljoin(base_url, link))                return absolute_links        async def get_or_scrape_url(self, url: str, max_cache_age: int = 86400) -> Dict[str, Any]:        """Get from cache or scrape if not available/expired"""        cache_key = self._get_cache_key(url)        cache_path = self.cache_dir / f"{cache_key}.json"                # Check if cache exists and is valid        if cache_path.exists():            try:                cache_data = json.loads(cache_path.read_text(encoding="utf-8"))                cached_at = cache_data.get("scraped_at", 0)                # Check if cache is still valid                if time.time() - cached_at < max_cache_age:                    return cache_data            except:                pass                # Cache miss or expired, perform scrape        result = await self.scrape_url(url)                # Save to cache if successful        if result["success"]:            self.cache_dir.mkdir(exist_ok=True)            with open(cache_path, "w", encoding="utf-8") as f:                json.dump(result, f, ensure_ascii=False, indent=2)                return result        def _get_cache_key(self, url: str) -> str:        """Generate a cache key from a URL"""        # Remove protocol and query parameters for consistent caching        normalized_url = re.sub(r'^https?://', '', url)        normalized_url = re.sub(r'\?.*$', '', normalized_url)        # Remove trailing slashes        normalized_url = normalized_url.rstrip('/')        # Create a hash to avoid file system issues with long URLs        url_hash = hashlib.md5(url.encode()).hexdigest()        # Return a file-system friendly key        return f"{url_hash}_{re.sub(r'[^a-zA-Z0-9_-]', '_', normalized_url)[:50]}"        def _configure_proxy(self, context) -> None:        """Configure proxy for playwright browser context"""        if not hasattr(self, 'proxy_settings') or not self.proxy_settings:            return                # Add proxy authentication if provided        auth = None        if self.proxy_settings.get('username') and self.proxy_settings.get('password'):            auth = {                'username': self.proxy_settings['username'],                'password': self.proxy_settings['password']            }                # Configure the proxy        proxy = {            'server': self.proxy_settings['server'],            'bypass': self.proxy_settings.get('bypass', '')        }        if auth:            proxy['username'] = auth['username']            proxy['password'] = auth['password']                # Apply proxy settings to the context        context.route('**/*', lambda route: route.continue_(proxy=proxy))        def extract_structured_data(self, html_content: str) -> Dict[str, Any]:        """Extract structured data from HTML page"""        try:            soup = BeautifulSoup(html_content, 'lxml')            structured_data = {}                        # Extract JSON-LD            json_ld_data = []            for script in soup.find_all('script', type='application/ld+json'):                try:                    data = json.loads(script.string)                    if isinstance(data, dict):                        json_ld_data.append(data)                    elif isinstance(data, list):                        json_ld_data.extend(data)                except Exception as e:                    logger.error(f"Error parsing JSON-LD: {e}")            if json_ld_data:                structured_data['json_ld'] = json_ld_data                        # Extract microdata            microdata = {}            for element in soup.find_all(itemscope=True):                if element.has_attr('itemtype'):                    item_type = element['itemtype']                    item_props = {}                                        # Extract properties                    for prop in element.find_all(itemprop=True):                        prop_name = prop['itemprop']                        # Get property value based on tag type                        if prop.name == 'meta':                            prop_value = prop.get('content', '')                        elif prop.name == 'img':                            prop_value = prop.get('src', '')                        elif prop.name == 'a':                            prop_value = prop.get('href', '')                        elif prop.name == 'time':                            prop_value = prop.get('datetime', prop.get_text())                        else:                            prop_value = prop.get_text().strip()                                                item_props[prop_name] = prop_value                                        if item_type not in microdata:                        microdata[item_type] = []                    microdata[item_type].append(item_props)            if microdata:                structured_data['microdata'] = microdata                        # Extract OpenGraph metadata            og_data = {}            for meta in soup.find_all('meta', property=re.compile(r'^og:')):                prop = meta.get('property', '').replace('og:', '')                content = meta.get('content', '')                if prop and content:                    og_data[prop] = content            if og_data:                structured_data['opengraph'] = og_data                        # Extract Twitter card metadata            twitter_data = {}            for meta in soup.find_all('meta', attrs={'name': re.compile(r'^twitter:')}):                prop = meta.get('name', '').replace('twitter:', '')                content = meta.get('content', '')                if prop and content:                    twitter_data[prop] = content            if twitter_data:                structured_data['twitter_card'] = twitter_data                        return structured_data        except Exception as e:            logger.error(f"Error extracting structured data: {e}")            return {}        async def _handle_rate_limiting(self, response):        """Handle rate limiting based on response codes"""        if response.status == 429:  # Too Many Requests            retry_after = response.headers.get('retry-after')            wait_time = int(retry_after) if retry_after and retry_after.isdigit() else 60            logger.info(f"Rate limited. Waiting for {wait_time} seconds")            await asyncio.sleep(wait_time)            return True                return False        async def capture_screenshot(self, url: str, full_page: bool = True) -> Dict[str, Any]:        """Capture screenshot of a webpage"""        try:            browser = await self.get_browser()            context = await browser.new_context(                user_agent=self.user_agent,                viewport={'width': 1920, 'height': 1080}            )            page = await context.new_page()            await page.goto(url, wait_until="networkidle")                        # Capture screenshot            screenshot_path = self.data_dir / "screenshots"            screenshot_path.mkdir(exist_ok=True)            filename = f"{hashlib.md5(url.encode()).hexdigest()}.png"            file_path = screenshot_path / filename            await page.screenshot(path=str(file_path), full_page=full_page)                        return {                "url": url,                "screenshot_path": str(file_path),                "timestamp": int(time.time()),                "success": True            }        except Exception as e:            logger.error(f"Screenshot failed for {url}: {e}")            return {                "url": url,                "error": str(e),                "success": False            }        async def scrape_with_pagination(self, url: str, max_pages: int = 5) -> Dict[str, Any]:        """Scrape a URL and follow pagination links"""        all_content = ""        current_url = url        pages_scraped = 0                while current_url and pages_scraped < max_pages:            result = await self.scrape_url(current_url)            if not result["success"]:                break                        # Accumulate content            all_content += result["content"] + "\n\n---\n\n"            pages_scraped += 1                        # Find next page link            next_url = self._find_next_page_link(result["content"], current_url)            if not next_url or next_url == current_url:                break                        current_url = next_url            # Add delay between pages            await asyncio.sleep(random.uniform(self.min_delay, self.max_delay))                # Create combined result        return {            "url": url,            "title": f"Paginated content ({pages_scraped} pages)",            "content": all_content,            "scraped_at": int(time.time()),            "success": True,            "pages_scraped": pages_scraped        }        def _find_next_page_link(self, content: str, current_url: str) -> Optional[str]:        """Find pagination link in content"""        soup = BeautifulSoup(content, "lxml")                # Common patterns for next page links        next_selectors = [            '.pagination .next',            '.pagination a[rel="next"]',            'a.next',            'a:contains("Next")',            'a[aria-label="Next"]',            '.pagination a:contains("›")',            '.pagination a:contains(">")'        ]                for selector in next_selectors:            try:                next_link = soup.select_one(selector)                if next_link and next_link.get('href'):                    return urljoin(current_url, next_link['href'])            except:                continue                return None        def _extract_main_content(self, soup: BeautifulSoup) -> str:        """Extract the main content from a page, ignoring navigation, ads, etc."""        # Remove common non-content elements        for element in soup.select('header, footer, nav, aside, .ads, .comments, .sidebar, script, style'):            element.decompose()                # Try to find the main content container        main_content = None        for selector in ['article', 'main', '.content', '.post', '#content', '[itemprop="articleBody"]']:            content = soup.select_one(selector)            if content:                main_content = content                break                if not main_content:            # Fall back to density-based content extraction if no container found            paragraphs_by_parent = {}            for p in soup.find_all('p'):                parent = p.parent                if parent not in paragraphs_by_parent:                    paragraphs_by_parent[parent] = []                paragraphs_by_parent[parent].append(p)            
            max_text_length = 0
            main_content = soup.body or soup
            for parent, paragraphs in paragraphs_by_parent.items():
                text_length = sum(len(p.get_text()) for p in paragraphs)
                if text_length > max_text_length:
                    main_content = parent
                    max_text_length = text_length
        
        return main_content.get_text(separator="\n").strip()